"""Scan Humanized .docx for common issues; print findings (no writes)."""
from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(doc):
    from docx.oxml.ns import qn

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _all_text(doc) -> list[tuple[str, str]]:
    """(where, text) for each paragraph."""
    out: list[tuple[str, str]] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if t.strip():
            out.append((f"p{i}", t))
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    t = p.text or ""
                    if t.strip():
                        out.append((f"t{ti}:r{ri}c{ci}:p{pi}", t))
    return out


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    path = root / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])

    doc = Document(str(path))
    chunks = _all_text(doc)
    big = "\n".join(t for _, t in chunks)

    checks = [
        ("et al., 20", r"et al\., 20\d\d"),
        ("(Author, A.", r"\([A-Z][a-z]+, [A-Z]\."),
        ("replacement char", "\ufffd"),
        ("  (double space)", "  "),
        ("Fast API", "Fast API"),
        ("Control A ", "Control A "),
        ("An compromised", "An compromised"),
        ("accordance to", "accordance to"),
        ("in five-window", "in five-window"),
        ("creates a prototype that end-to-end", "creates a prototype that end-to-end"),
        ("F1 = 99.86 and 99.42", "F1 = 99.86 and 99.42"),
    ]
    print(path.name, "total chars", len(big))
    for label, pat in checks:
        if pat.startswith("(") and len(pat) < 40:
            n = len(re.findall(pat, big))
        elif "\\" in pat:
            n = len(re.findall(pat, big))
        else:
            n = big.count(pat)
        if n:
            print(f"  [{label}] count={n}")

    # suspicious words
    sus = ["utilized", "in order to", "leverage", "robust", "delve"]
    for w in sus:
        c = len(re.findall(rf"\b{re.escape(w)}\b", big, re.I))
        if c:
            print(f"  [word: {w}] count={c}")


if __name__ == "__main__":
    main()
