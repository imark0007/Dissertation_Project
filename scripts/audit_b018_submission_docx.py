"""Audit B01821011_Arka_Talukder_Dissertation_Final.docx (read-only)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "submission" / "B01821011_Arka_Talukder_Dissertation_Final.docx"


def all_doc_text(doc: Document) -> str:
    """Paragraphs and table cell text, orderApproximate (good enough for audit)."""
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text or ""
        if t.strip():
            parts.append(t.strip())
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
    return "\n".join(parts)


def word_count(s: str) -> int:
    return len(re.findall(r"[A-Za-z0-9']+", s))


def contiguous(nums: list[int], start: int) -> bool:
    if not nums:
        return False
    s = sorted(set(nums))
    return s == list(range(start, s[-1] + 1))


def main() -> None:
    doc = Document(str(DOCX))
    text = all_doc_text(doc)
    low = text.lower()

    print("=" * 60)
    print("AUDIT:", DOCX.name)
    print("Total words (approx):", word_count(text))
    print("Em-dash (U+2014):", text.count("\u2014"))
    print("=" * 60)

    # Chapters: paragraph starts
    ch_nums: set[int] = set()
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        m = re.match(r"^Chapter (\d+)", t, re.I)
        if m and len(t) < 200:
            ch_nums.add(int(m.group(1)))
    print("Chapter headings (short lines) found:", sorted(ch_nums))
    missing_ch = [n for n in range(1, 14) if n not in ch_nums]
    if missing_ch:
        print("  WARNING — missing short chapter headings for:", missing_ch)
    else:
        print("  Chapters 1-13: OK (heading scan)")

    # Front / forms
    print("\n--- Front / administrative ---")
    items = [
        ("B01821011 / Arka", "b01821011" in low or "arka talukder" in low),
        ("Abstract", re.search(r"\babstract\b", low) is not None),
        ("Table of contents", "table of contents" in low),
        ("List of figures", "list of figures" in low),
        ("List of tables", "list of tables" in low),
        ("List of Abbreviations", "list of abbrev" in low or "abbreviation" in low),
        ("Acknowledgements", "acknowledg" in low),
        (
            "Declaration of originality",
            re.search(r"declaration", low) and re.search(r"originality|own work|plagi", low) is not None,
        ),
        (
            "Library / ethics (dataset)",
            re.search(
                r"library|release form|ethics|non-sensitive|public data|not require.*ethics|consent", low
            )
            is not None,
        ),
    ]
    for name, ok in items:
        print(f"  {name}: {'OK' if ok else 'CHECK'}")

    print("\n--- Appendices ---")
    for ap in "ABCD":
        ok = f"appendix {ap.lower()}" in low
        if not ok:
            ok = re.search(rf"appendix\s+{ap}\b", low) is not None
        print(f"  Appendix {ap}: {'OK' if ok else 'MISSING'}")

    # Main RQ & project-spec alignment (exact strings from MD audit script)
    main_rq = (
        "How can an explainable dynamic graph neural network, trained using federated learning, "
        "detect attacks in Software-Defined IoT flow data and generate SIEM alerts that support SOC "
        "operations on CPU-based edge devices?"
    )
    print("\n--- Research questions ---")
    print("  Main RQ exact (as in audit script):", "YES" if main_rq in text else "NO (paraphrased)")

    # Sub-Q wording variants
    s1a = "Does a dynamic graph model perform better" in text
    s1b = "outperform" in text and "sub" in low
    s2a = "Can federated learning maintain similar performance" in text
    s2b = "federated learning" in low and "raw" in low
    s3a = "Can the model generate useful explanations" in text
    s3b = "explanation" in low and "triage" in low
    print("  Sub-Q1 (legacy checklist text):", "YES" if s1a else "N/A/variant (see outperf)", s1b)
    print("  Sub-Q2 (legacy checklist text):", "YES" if s2a else "N/A/variant", s2a or s2b)
    print("  Sub-Q3 (legacy checklist text):", "YES" if s3a else "N/A/variant", s3a or s3b)

    # Marking / criteria table
    has_spec_table = "introduction" in low and "context" in low and "literature" in low
    has_weight = "5%" in text or "10%" in text or "20%" in text or "25%" in text
    print("  Criterion/weighting language:", "OK" if (has_spec_table and has_weight) else "CHECK")

    # Figure / table numbers in narrative
    fig_nums = [int(n) for n in re.findall(r"Figure\s+(\d+)", text, flags=re.I)]
    tab_nums = [int(n) for n in re.findall(r"Table\s+(\d+)", text, flags=re.I)]
    if fig_nums:
        mx = max(fig_nums)
        print("\n--- Figure / Table numbering (text scan) ---")
        print("  Figure refs min..max:", min(fig_nums), "..", mx)
        gaps = [i for i in range(1, mx + 1) if i not in set(fig_nums)]
        print("  Gaps in 1..max (0 means contiguous refs):", len(gaps), ("..." + str(gaps[:20]) if gaps else ""))
        print("  Contiguous 1..N in refs:", contiguous(fig_nums, 1))
    if tab_nums:
        mxt = max(tab_nums)
        print("  Table refs min..max:", min(tab_nums), "..", mxt)
        g2 = [i for i in range(1, mxt + 1) if i not in set(tab_nums)]
        print("  Gaps in 1..max tables:", g2 if g2 else "none")
    a1 = len(re.findall(r"(?:fig(?:ure)?\s+a1[-]?\d+)|a1[-]?\d+", text, re.I))
    print("  Appendix A1 figure refs (rough):", a1)

    # Technical completeness
    print("\n--- Technical / evidence ---")
    tech = {
        "CICIoT2023 / Pinto": "ciciot" in low or "pinto" in low,
        "GAT + GRU / dynamic GNN": "gat" in low and "gru" in low,
        "Flower FedAvg": "flower" in low and "fedavg" in low,
        "Captum / Integrated Gradients": "captum" in low or "integrated" in low,
        "FastAPI / ECS / SIEM": ("fastapi" in low or "ecs" in low) and "siem" in low,
        "CPU / edge": "cpu" in low and "edge" in low,
        "Confusion / ROC (evaluation)": "confusion" in low or "roc" in low,
        "Future work (Ch 9)": "future work" in low or "recommendation" in low,
        "Limitations (honesty)": "limitation" in low,
        "Critical self-evaluation (Ch 10)": re.search(r"critical self", low) is not None,
    }
    for k, v in tech.items():
        print(f"  {k}: {'OK' if v else 'CHECK'}")

    # Placeholder / TODO
    ph = ["[insert", "todo", "xxx", "TBD", "lorem", "UWS LOGO (insert official logo here)"]
    hits = [x for x in ph if x in low]
    print("\n--- Placeholders / TODO ---")
    print("  None" if not hits else str(hits))

    print("\n" + "=" * 60)
    print("This script does not prove Word fields (TOC page numbers) are updated.")
    print("=" * 60)


if __name__ == "__main__":
    main()
