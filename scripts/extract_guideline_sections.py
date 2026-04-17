"""
Extract paragraphs/headings from 'Updated_Guideline For FINAL REPORT.docx'.

Usage:
    python scripts/extract_guideline_sections.py
"""
from __future__ import annotations

from pathlib import Path


def main() -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise SystemExit("python-docx is required: pip install python-docx") from e

    p = Path("Updated_Guideline For FINAL REPORT.docx")
    if not p.exists():
        raise SystemExit(f"Missing: {p}")

    d = Document(str(p))
    print(f"Paragraphs: {len(d.paragraphs)}\n")
    for i, para in enumerate(d.paragraphs):
        t = para.text.strip()
        if not t:
            continue
        style = para.style.name if para.style else ""
        print(f"{i:04d}|{style}|{t}")


if __name__ == "__main__":
    main()

