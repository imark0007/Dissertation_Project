"""
Fix clear errors in Humanized submission .docx so facts and house style align with the
Final_Submission docx / handbook, without rewriting prose.

- Grammar / typo fixes found by comparison
- Missing % on F1 values in abstract results sentence
- (Author et al., Year) -> (Author et al. Year) for in-text Harvard
- Strip replacement-character bullets
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    """Yield paragraphs and tables in document order (body only)."""
    from docx.oxml.ns import qn
    from docx.table import _Cell

    if parent is None:
        return
    for child in parent.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _fix_text(s: str, *, citations_only: bool = False) -> str:
    if not s:
        return s
    if citations_only:
        pairs = []
    else:
        pairs = [
        ("Control A controlled", "A controlled"),
        ("An compromised", "A compromised"),
        ("Fast API", "FastAPI"),
        ("1.3 Research Aim and Question.", "1.3 Research Aim and Questions"),
        ("approachable sub-set", "manageable subset"),
        ("Federation of Flower FedAvg", "Flower FedAvg federation"),
        ("in accordance to the UWS", "in accordance with the UWS"),
        ("in accordance to ", "in accordance with "),
        (
            "On average, the inference of CPU is approximately 23 ms in five-window sequence, and federated communication is approximately 31 MB in ten rounds.",
            "Mean CPU inference is about 23 ms per five-window sequence, and federated communication is about 31 MB over ten rounds.",
        ),
        # Abstract (second paragraph): minimal grammar / clarity; keeps humanized wording elsewhere
        (
            "This dissertation creates a prototype that end-to-end on CICIoT2023",
            "This dissertation presents an end-to-end prototype on CICIoT2023",
        ),
        (
            "The architecture is trained on three non-IID clients (Dirichlet alpha = 0.5) centrally and using Flower FedAvg.",
            "The same architecture is trained centrally and with Flower FedAvg across three non-IID clients (Dirichlet alpha = 0.5).",
        ),
        (
            "The explainability model integrates Captum Integrated Gradients and GAT attention and outputs results in the form of ECS-like JSON as a FastAPI endpoint.",
            "Explainability combines Captum Integrated Gradients and GAT attention, and results are served as ECS-like JSON through a FastAPI endpoint.",
        ),
        ("convergence and accuracy relies on", "convergence and accuracy rely on"),
        # Align wording / facts with Final_Submission
        (
            "designed to facilitate such research",
            "made to support such research",
        ),
        (
            "it utilizes that public release in such a way",
            "it uses that public release in such a way",
        ),
        (" and can utilize relational", " and can use relational"),
        (
            "The reproducibility is facilitated by the use of a publicly available dataset",
            "The work is reproducible thanks to a publicly available dataset",
        ),
        (
            "how can a prototype system, moreover, use an explainable dynamic graph neural network",
            "how can a prototype system use an explainable dynamic graph neural network",
        ),
        (
            "They proved that with feature similarity",
            "They showed that with feature similarity",
        ),
        (
            "Training of the centralised Dynamic GNN took 6 epochs (with early-stopping set in the training script). Table 4 includes train loss and validation F1 / ROC-AUC in an epoch (results/metrics/gnn_training_history.json). ROC-AUC = 100.0% since epoch 1 in this run; train loss decreases to 0.0001 by epoch 4 and remains there until epoch 6.",
            "The centralised Dynamic GNN was trained for 6 epochs (with early-stopping configured in the training script). Table 4 lists train loss and validation F1 / ROC-AUC per epoch (results/metrics/gnn_training_history.json). Validation F1 = ROC-AUC = 100.0% from epoch 1 onward on this run; train loss falls from 0.484 to 0.0001 by epoch 4 and stays there through epoch 6.",
        ),
        (
            "Each record is in the shape of the ECS: event metadata, rule name, threat indicator, ML prediction and score and explanation (top_features, top_nodes). Whether these areas are adequate to SOC triage is addressed in Chapter 9.",
            "Each record follows the ECS-like shape: event metadata, rule name, threat indicator, ML prediction and score, and explanation (top_features, top_nodes). Whether these fields are sufficient for SOC triage is discussed in Chapter 9.",
        ),
        (
            "Serve the model with a FastAPI endpoint which provides ECS-like JSON alerts",
            "Serve the model through a FastAPI endpoint that returns ECS-like JSON alerts",
        ),
        # Section headings (Humanized typos / drift → Final)
        ("3.6 Interim Report Feedback Inc.", "3.6 Interim Report Feedback Incorporated"),
        (
            "6.10 Implementation Code Photos (Author Userbase)",
            "6.10 Implementation Code Screenshots (Author’s Codebase)",
        ),
        (
            "4.3.1 Data location (repository Layout)",
            "4.3.1 Where the Data Lives (repository Layout)",
        ),
        (
            "5.2 Pipeline, Alerts and Deployment (Conceptual)",
            "5.2 Pipeline, Alerts, and Deployment (Conceptual)",
        ),
        ("6.3 Data loading and Preprocessing.", "6.3 Data Loading and Preprocessing"),
        (
            "6.9 FastAPI Deployment and inference on CPU.",
            "6.9 FastAPI Deployment and CPU Inference",
        ),
        (
            "7.6 Statistics of Data and Experiment.",
            "7.6 Dataset and Experiment Statistics",
        ),
        (
            "8.10 Comparison with the previous work on CICIoT2023.",
            "8.10 Comparison with Prior Work on CICIoT2023",
        ),
        (
            "9.2 Organized Conclusion (Programme Format)",
            "9.2 Structured Conclusion (Programme Format)",
        ),
        ("9.5 Use University and Course Materials.", "9.5 Use of University and Course Materials"),
        (
            "10.3 Literature and Response to the Questions.",
            "10.3 Literature and Alignment with the Questions",
        ),
        (
            "10.4 Implementation: It Was worse than it Seemed.",
            "10.4 Implementation: What Was Harder Than It Looked",
        ),
        (
            "10.5 Results, Honesty and the 100% Question.",
            "10.5 Results, Honesty, and the “100%” Question",
        ),
        (
            "10.7 Time Management: My Reorders.",
            "10.7 Time Management: What I Would Reorder",
        ),
        # Chapter 2 numbering (SIEM must be §2.4 after §2.3 IoT)
        (
            "2.3 SIEM, SOC Workflows, and Alert Quality",
            "2.4 SIEM, SOC Workflows, and Alert Quality",
        ),
        (
            "2.5 Dynamic Graphs and Graph Neural Networks.",
            "2.5 Graph Neural Networks and Dynamic Graphs",
        ),
        ]
    for a, b in pairs:
        s = s.replace(a, b)
    if not citations_only:
        s = re.sub(r"\bUtilized\b", "Used", s)
        s = re.sub(r"\butilized\b", "used", s)
        s = re.sub(r" in order to ", " to ", s)
        if s.startswith("In order to "):
            s = "To " + s[12:]
    if not citations_only:
        # Abstract-style F1 percentages (match Final_Submission)
        s = re.sub(
            r"F1 = 99\.86 and 99\.42\.",
            "F1 = 99.86% and 99.42%.",
            s,
        )
    # In-text Harvard: no comma before year inside ( ... )
    s = re.sub(
        r"\(([A-Z][A-Za-zÀ-ÿ'\-]+ et al\.), (\d{4})\)",
        r"(\1 \2)",
        s,
    )
    s = re.sub(
        r"\(([A-Z][A-Za-zÀ-ÿ'\-]+ and [A-Z][A-Za-zÀ-ÿ'\-]+), (\d{4})\)",
        r"(\1 \2)",
        s,
    )
    # "(e.g. Author et al., 2020)" -> comma removed before year
    s = re.sub(
        r"\(e\.g\. ([A-Z][A-Za-zÀ-ÿ'\-]+ et al\.), (\d{4})\)",
        r"(e.g. \1 \2)",
        s,
    )
    # "see Author et al., 2023" -> parenthetical form
    s = re.sub(
        r"see ([A-Z][A-Za-zÀ-ÿ'\-]+ et al\.), (\d{4})",
        r"see (\1 \2)",
        s,
    )
    if not citations_only:
        s = s.replace("\ufffd", "•")
        # Collapse accidental double spaces (preserve single spaces only)
        s = re.sub(r" {2,}", " ", s)
    return s


def _fix_paragraph(p: Paragraph, *, citations_only: bool = False) -> int:
    t = p.text
    if not t:
        return 0
    new = _fix_text(t, citations_only=citations_only)
    if new == t:
        return 0
    if not p.runs:
        p.add_run(new)
    else:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    return 1


def _fix_table(tbl: Table, *, citations_only: bool = False) -> int:
    n = 0
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                n += _fix_paragraph(p, citations_only=citations_only)
    return n


def _sync_abstract_from_final(hum: Document, final: Document) -> None:
    """Replace the four abstract body paragraphs in Humanized with Final text (wording + ~188 words)."""
    try:
        idx_h = next(i for i, p in enumerate(hum.paragraphs) if p.text.strip() == "1. Abstract")
        idx_f = next(i for i, p in enumerate(final.paragraphs) if p.text.strip() == "1. Abstract")
    except StopIteration:
        return
    for j in range(4):
        hum.paragraphs[idx_h + 1 + j].text = final.paragraphs[idx_f + 1 + j].text


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    path = root / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"
    final_path = root / "submission" / "Arka_Talukder_Dissertation_Final_Submission.docx"
    citations_only = False
    args = [a for a in sys.argv[1:] if not a.startswith("-")]
    flags = [a for a in sys.argv[1:] if a.startswith("-")]
    if args:
        path = Path(args[0])
    if "--citations-only" in flags:
        citations_only = True

    doc = Document(str(path))
    if final_path.is_file():
        _sync_abstract_from_final(doc, Document(str(final_path)))

    changed = 0
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            changed += _fix_paragraph(block, citations_only=citations_only)
        else:
            changed += _fix_table(block, citations_only=citations_only)

    # Corrupted style name from some exports (empty paragraphs)
    try:
        normal = doc.styles["Normal"]
        for p in doc.paragraphs:
            if p.style and p.style.name == "age":
                p.style = normal
    except KeyError:
        pass

    out = path
    doc.save(str(out))
    print(f"Saved {out}; blocks updated: {changed}")


if __name__ == "__main__":
    main()
