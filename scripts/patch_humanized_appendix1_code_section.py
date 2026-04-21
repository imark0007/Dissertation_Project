"""One-off: refresh Handbook Appendix 1 code-figure prose in Humanized .docx only."""
from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
PATH = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"

# En dash for ranges (Word-friendly)
ND = "\u2013"


def main() -> None:
    doc = Document(str(PATH))
    try:
        start = next(
            i
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip().startswith("Appendix A: Handbook Appendix 1")
        )
    except StopIteration:
        raise SystemExit("Could not find Appendix A: Handbook Appendix 1 heading") from None

    # Expected fixed block: heading + 22 body lines (intro 4 + 6*(figure+caption+interpret))
    texts = [
        "Handbook Appendix 1 code views must be provided as figures, each with a caption and a description of how to interpret the snippet within this dissertation. "
        f"Figure A1-1{ND}Figure A1-6 below meet that requirement.",
        "The submission codebase generates the bitmaps so line numbers stay aligned with the files on disk. After editing source, regenerate with:",
        "python scripts/render_appendix1_code_figures.py",
        "Output directory: results/figures/appendix1/. These appear in the Table of Figures as "
        f"Figure A1-1{ND}Figure A1-6 (appendix labels), independent of the main-body sequence Figure 1{ND}Figure 29. "
        "Chapter 6 (Section 6.10) and this appendix both use dark-theme, line-numbered core snippets "
        "(not full-file captures) to help examiners map prose to code.",
        f"Figure A1-1 - Dynamic graph classifier: DynamicGNN sequence forward core (GRU vs. mean-pool ablation, logits). (Chapter 13, Appendix C.) Source: src/models/dynamic_gnn.py, lines 75{ND}97.",
        "Caption (formal): Figure A1-1 - Core implementation of the dynamic GNN (DynamicGNN): node embedding, two GATConv layers with multi-head attention and dropout, per-window graph embedding, sequence encoding with GRU (or mean-pool when use_gru is false for ablation in Section 8.7), and two-class logits. Attention weights can be kept for explainability (return_attention_weights).",
        "Interpretation: This is the primary learnable model in Chapters 5–6. Each time step is a PyTorch Geometric Data object (nodes = flows in a window, edges = kNN). _encode_graph runs GAT message passing and pools node states to one vector per window; forward walks the window sequence in time and either uses the GRU (full model) or mean-pools across time (ablation in Section 8.7). from_config binds hyperparameters to config/experiment.yaml, which supports the sensitivity study in Section 8.8.",
        f"Figure A1-2 - Building a single-window kNN graph from rows of flow features. (Chapter 13, Appendix C.) Source: src/data/graph_builder.py, lines 37{ND}58.",
        "Caption (formal): Figure A1-2 - flows_to_knn_graph: for one window of N flows with F features, fits sklearn.neighbors.NearestNeighbors (Euclidean), adds bidirectional edges between each node and its k nearest neighbours (capped when N is small), and returns a torch_geometric.data.Data with x, edge_index, and graph-level label y.",
        "Interpretation: This is how graph structure is defined without device IPs (Chapter 1): similarity in the 46-dimensional flow-feature space stands in for physical topology. Bidirectional edges suit GAT. The caller supplies the graph label (benign vs. attack); that matches the stratified pool in Figure A1-3, not a per-flow vote—important when reading the evaluation chapters.",
        f"Figure A1-3 - Stratified windowing so both classes appear in training graphs. (Chapter 13, Appendix C.) Source: src/data/graph_builder.py, lines 106{ND}123.",
        "Caption (formal): Figure A1-3 - build_graphs_for_split: splits flows into benign and attack pools, builds sliding (or strided) windows from each pool via _build_windows_from_pool, balances class counts, shuffles, and logs totals—addressing severe imbalance in raw CICIoT2023.",
        "Interpretation: This explains why training does not collapse to always predict attack despite a very high attack rate in the raw CSV (Chapter 4): windows are drawn within each class, so each graph's label matches its pool. minority_stride increases overlap for the smaller pool when needed. Window size and k for the dissertation come from config/experiment.yaml and feed the sensitivity grid in Section 8.8.",
        f"Figure A1-4 - Post-hoc explanations: forward pass with attention, Integrated Gradients, top nodes/features. (Chapter 13, Appendix C.) Source: src/explain/explainer.py, lines 53{ND}95.",
        "Caption (formal): Figure A1-4 - explain_sequence: runs the model with attention enabled, wraps Captum Integrated Gradients on the last window's node features (_ig_wrapper), sums absolute attributions to rank top nodes and top feature indices, and returns an ExplanationBundle for ECS-like JSON alerts.",
        "Interpretation: This links Chapter 6 to Chapter 8 example alerts: SOC-facing explanations follow IG magnitudes and GAT attention from the same forward pass used at inference. IG is applied on the final graph in the sequence (design choice in code comments). If Captum is unavailable, behaviour degrades gracefully (HAS_CAPTUM in the same module).",
        f"Figure A1-5 - Federated learning CLI: Flower server vs. client, local data, GNNFlowerClient. (Chapter 13, Appendix C.) Source: src/federated/run_federated.py, lines 28{ND}71.",
        "Caption (formal): Figure A1-5 - main in run_federated.py: loads YAML config; server mode calls run_fl_server with num_rounds and minimum clients from fl in the config; client mode loads data/graphs/client_{cid}_graphs.pt, builds sequence loaders, instantiates GNNFlowerClient, and calls fl.client.start_numpy_client to 127.0.0.1:8080.",
        "Interpretation: This is the entry point for Chapter 8 federated runs: each client trains only on its partition (non-IID split from src.federated.data_split). Rounds and client settings come from config/experiment.yaml under fl. The flow matches FedAvg in Chapter 2—local training, then server aggregation (src/federated/server.py, not shown).",
        f"Figure A1-6 - HTTP API: POST /score builds graphs, runs explain_sequence, emits ECS-like alert JSON. (Chapter 13, Appendix C.) Source: src/siem/api.py, lines 67{ND}89.",
        "Caption (formal): Figure A1-6 - FastAPI: startup loads weights and config; POST /score accepts flow windows, builds kNN graphs with knn_k from config via flows_to_knn_graph, calls explain_sequence (top-5 nodes and features), records wall-clock milliseconds, and returns format_ecs_alert together with prediction and score.",
        "Interpretation: This is the deployment surface from Chapters 1 and 5: one HTTP call turns raw feature windows into SIEM-shaped JSON for triage. Inference k matches config/experiment.yaml, limiting train/serve skew. Chapter 8 latency figures refer to this synchronous CPU path.",
    ]

    for offset, text in enumerate(texts):
        doc.paragraphs[start + 1 + offset].text = text

    doc.save(str(PATH))
    print("Patched", PATH.name, "paragraphs", start + 1, "..", start + len(texts))


if __name__ == "__main__":
    main()
