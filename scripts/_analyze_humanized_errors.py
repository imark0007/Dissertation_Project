"""Read-only deep scan of Humanized .docx: report likely errors (no writes)."""
from __future__ import annotations

import collections
import re
import zipfile
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
PATH = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"

CHAP = re.compile(r"^Chapter\s+(\d+)\s+[–-]", re.I)
FIG_CAP = re.compile(r"^Figure\s+(\d+)\s*:", re.I)
FIG_A1 = re.compile(r"^Figure\s+A1-\s*(\d+)\s*:", re.I)
SUBSEC = re.compile(
    r"^(\d{1,2})\.(\d{1,2})(?:\.(\d{1,2}))?\s+"
)  # 1.3 or 4.3.1 at line start

FORBIDDEN = re.compile(
    r"\b(?:delve|crucial|comprehensive|leverage|utilize|facilitate|myriad|plethora|"
    r"notably|underscores|pivotal|groundbreaking|seamlessly|tapestry|landscape|"
    r"nuanced\s+understanding|moving\s+forward|it\s+is\s+important\s+to\s+note\s+that|"
    r"this\s+highlights|in\s+light\s+of|overarching|paradigm|robust(?:ly)?)\b",
    re.I,
)
OVERCLAIM = re.compile(
    r"\b(?:prove[sd]?|proven|guaranteed|definitive|unquestionably|eliminates|"
    r"zero-risk|best\s+overall|perfect(?:ly)?\s+accurate)\b",
    re.I,
)
GRAMMAR_SUS = [
    (r"\bAn\s+(?:compromised|attack|IoT)\b", "article/determiner (An + consonant sound)"),
    (r"\baccordance\s+to\b", "usually 'in accordance with'"),
    (r"\bIt\s+is\s+not\s+uncommon\s+situations\b", "grammar ('situations' awkward)"),
    (r"\bSIEM\s+and\s+other\s+tools\s+like\s+them\s+allow\s+SIEM\b", "redundant wording"),
    (r"\bportion\s+of\s+devices\s+is\b", "grammar (portion vs devices agreement)"),
    (r"\bremain(?:s)?\s+a\s+under-explored\b", "use 'an under-explored'"),
    (r"\bScalability\s+Manual\s+inspection\b", "likely missing punctuation/newline"),
    (r"\bControl\s+A\b", "placeholder? (Control A)"),
    (r"\bet\s+al\s+(?!\.|\()", "et al spacing before year"),
    (r"\b(?:very\s+){2,}", "double intensifier"),
    (r"[ \t]{2,}", "multiple spaces in paragraph (per line)"),
]
TYPO_SUS = [
    ("beachhead", None),  # valid word - skip
    (r"\bform\s+of\s+temporal\s+patterns\s+of\s+communication\b", None),
]


def all_paragraph_texts(doc: docx.document.Document) -> list[str]:
    out: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    t = (cp.text or "").strip()
                    if t:
                        out.append(t)
    return out


def main_body_paragraphs(doc: docx.document.Document) -> list[str]:
    paras = [((p.text or "").strip()) for p in doc.paragraphs if (p.text or "").strip()]
    idx = None
    for i, t in enumerate(paras):
        m = CHAP.match(t)
        if m and m.group(1) == "1":
            idx = i
            break
    if idx is None:
        return paras
    end = len(paras)
    for j in range(idx + 1, len(paras)):
        m = CHAP.match(paras[j])
        if m and m.group(1) == "11":
            end = j
            break
    return paras[idx:end]


def analyze_figure_captions(paras: list[str]) -> dict:
    nums = []
    for t in paras:
        m = FIG_CAP.match(t)
        if m:
            nums.append(int(m.group(1)))
    a1 = []
    for t in paras:
        m = FIG_A1.match(t)
        if m:
            a1.append(int(m.group(1)))
    issues = []
    if nums:
        for i in range(len(nums) - 1):
            if nums[i + 1] not in (nums[i], nums[i] + 1, nums[i] + 2):
                # allow duplicate refs in different contexts
                if nums[i + 1] < nums[i]:
                    issues.append(f"Figure number decreases: {nums[i]} → {nums[i + 1]}")
        mx, mn = max(nums), min(nums)
        if mn != 1:
            issues.append(f"First body Figure number is {mn}, not 1")
        exp = list(range(mn, mx + 1))
        missing = [n for n in exp if n not in nums]
        dupes = [n for n, c in collections.Counter(nums).items() if c > 1]
        if missing:
            issues.append(f"Possible missing Figure nums in sequence (seen set gaps): {missing[:15]}")
        if dupes:
            issues.append(f"Duplicate Figure labels in captions: {sorted(dupes)[:20]}")
    return {"figure_nums_sample": nums[:5], "figure_nums_max": max(nums) if nums else 0, "appendix_a1": a1, "issues": issues}


def subsection_heading_lines(doc: docx.document.Document) -> list[tuple[str, str]]:
    """(style_name, text) for paras that look like numbered section headings."""
    out = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        st = p.style.name if p.style else ""
        if SUBSEC.match(t) and (
            st.startswith("Heading") or st == "Title" or (st == "Normal" and re.match(r"^\d+\.\d+\s", t))
        ):
            out.append((st, t[:100]))
    return out


def main() -> None:
    doc = docx.Document(str(PATH))
    paras = all_paragraph_texts(doc)
    body = main_body_paragraphs(doc)
    big = "\n".join(paras)

    print("=== File ===")
    print(PATH.name)
    print("Non-empty paragraph-like chunks:", len(paras))
    print("Main body (Ch1–before Ch11) paragraphs:", len(body))
    print()

    print("=== OOXML: unicode replacement / private-use ===")
    z = zipfile.ZipFile(PATH)
    raw = z.read("word/document.xml").decode("utf-8")
    rep = raw.count("\ufffd")
    print("U+FFFD in XML:", rep)
    print()

    print("=== Heuristic phrase flags (not proof of error) ===")
    for pat, label in (
        (FORBIDDEN, "filler-ish / AI-shaped"),
        (OVERCLAIM, "overclaim / absolutist"),
    ):
        hits = list(pat.finditer(big))
        print(f"{label}: {len(hits)} hits")
        for m in hits[:8]:
            s = max(0, m.start() - 50)
            e = min(len(big), m.end() + 50)
            snippet = big[s:e].replace("\n", " ")
            print(f"  … {snippet[:160]} …")
    print()

    print("=== Grammar / wording suspects (regex) ===")
    for pat, note in GRAMMAR_SUS:
        for i, t in enumerate(paras):
            if re.search(pat, t):
                safe = (t[:180] + ("..." if len(t) > 180 else "")).encode("ascii", "replace").decode("ascii")
                print(f"  [{note}] para chunk ~{i}: {safe}")
    print()

    print("=== audit_humanized_docx-style spot checks ===")
    for label, pat in [
        ("replacement char in joined text", "\ufffd"),
        ("  double-space", "  "),
    ]:
        if pat == "\ufffd":
            n = big.count(pat)
        else:
            n = sum(t.count("  ") for t in paras)
        if n:
            print(f"  {label}: {n}")
    print()

    print("=== Figure captions (body paras) ===")
    figinfo = analyze_figure_captions([(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()])
    for k, v in figinfo.items():
        print(f"  {k}: {v}")
    print()

    print("=== Numbered headings: Normal style (possible outline drift) ===")
    subs = subsection_heading_lines(doc)
    normal_num = [x for x in subs if x[0] == "Normal" and re.match(r"^\d+\.\d+", x[1])]
    print(f"  Count: {len(normal_num)} (show up to 12)")
    for st, tx in normal_num[:12]:
        print(f"    [{st}] {tx}")
    print()

    print("=== Same subsection number, multiple Heading styles sample ===")
    by_num: dict[str, list[str]] = collections.defaultdict(list)
    for st, tx in subs:
        m = SUBSEC.match(tx)
        if m:
            key = ".".join(g for g in m.groups() if g is not None)
            by_num[key].append(st)
    weird = {k: v for k, v in by_num.items() if len(set(v)) > 1}
    for k in sorted(weird, key=lambda x: tuple(map(int, x.split("."))))[:15]:
        print(f"  {k}: styles {set(by_num[k])}")


if __name__ == "__main__":
    main()
