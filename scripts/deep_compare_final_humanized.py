"""Symmetric deep scan: AI-ish phrases, claims, formatting hints — both docx."""
from __future__ import annotations

import collections
import re
import zipfile
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
PATH_FINAL = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission.docx"
PATH_HUM = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"

CHAP = re.compile(r"^Chapter\s+(\d+)\s+[–-]")

# Coach list + extra AI-ish academic fillers (heuristic only)
FORBIDDEN = re.compile(
    r"\b(?:delve|crucial|comprehensive|leverage|utilize|facilitate|myriad|plethora|"
    r"notably|underscores|pivotal|groundbreaking|seamlessly|tapestry|landscape|"
    r"nuanced\s+understanding|moving\s+forward|it\s+is\s+important\s+to\s+note\s+that|"
    r"this\s+highlights|in\s+light\s+of|overarching|paradigm|robust(?:ly)?)\b",
    re.I,
)
OVERCLAIM = re.compile(
    r"\b(?:prove[sd]?|proven|guaranteed|definitive|unquestionably|eliminates|"
    r"zero-risk|best\s+overall|perfect(?:ly)?\s+accurate)\b",
    re.I,
)
STRUCTURAL_AI = re.compile(
    r"\b(?:Furthermore|Moreover|Additionally|In\s+conclusion|In\s+summary|"
    r"To\s+summarize|It\s+is\s+worth\s+noting|plays\s+a\s+(?:crucial|vital|key)\s+role|"
    r"serves\s+as\s+a\s+(?:foundation|framework))\b",
    re.I,
)

US_VS_UK = [
    (r"\borganize\b", "US organize"),
    (r"\banalyze\b", "US analyze (if not noun 'analysis' context)"),
    (r"\bcolor\b", "US color"),
    (r"\bbehavior\b", "US behavior"),
    (r"\brecognize\b", "US recognize"),
    (r"\bfavor\b", "US favor"),
    (r"\boptimize\b", "US optimize"),
]


def paras(path: Path) -> list[str]:
    return [p.text.strip() for p in docx.Document(str(path)).paragraphs if p.text.strip()]


def main_body_text(paras: list[str]) -> str:
    idx = {int(m.group(1)): i for i, t in enumerate(paras) if (m := CHAP.match(t))}
    if 1 not in idx or 11 not in idx:
        return "\n".join(paras)
    return "\n".join(paras[idx[1] : idx[11]])


def count_pattern(pat: re.Pattern, text: str) -> tuple[int, list[str]]:
    hits = pat.findall(text)
    samples = []
    for m in pat.finditer(text):
        start = max(0, m.start() - 40)
        end = min(len(text), m.end() + 40)
        samples.append(text[start:end].replace("\n", " "))
        if len(samples) >= 8:
            break
    return len(hits), samples


def ooxml_stats(path: Path) -> dict:
    out: dict = {}
    with zipfile.ZipFile(path) as z:
        doc = z.read("word/document.xml").decode("utf-8")
    out["paragraph_markers"] = len(re.findall(r"<w:p\b", doc))
    out["pStyle_Normal"] = len(re.findall(r'<w:pStyle w:val="Normal"', doc))
    out["pStyle_Heading"] = len(re.findall(r'<w:pStyle w:val="Heading', doc))
    out["w14_w15_in_doc"] = bool(re.search(r"<w14:|w15:", doc))
    # Line spacing: w:line in twips (common 360 = 1.5 lines at 240 twips/line default)
    lines = re.findall(r'<w:spacing[^>]*w:line="(\d+)"', doc)
    out["spacing_line_vals_top"] = collections.Counter(lines).most_common(6)
    # Font size half-points w:sz w:val=
    sz = re.findall(r'<w:sz w:val="(\d+)"', doc)
    out["sz_half_pts_top"] = collections.Counter(sz).most_common(8)
    # Section breaks
    out["sectPr"] = len(re.findall(r"<w:sectPr", doc))
    out["page_break"] = len(re.findall(r"w:br w:type=\"page\"", doc))
    # Hyphenation / widow control tags present
    out["hyphenation_off_tags"] = doc.count("w:suppressAutoHyphens")
    return out


def style_names_sample(path: Path, limit: int = 800) -> collections.Counter:
    c: collections.Counter = collections.Counter()
    d = docx.Document(str(path))
    for i, p in enumerate(d.paragraphs):
        if i >= limit:
            break
        name = p.style.name if p.style else ""
        c[name] += 1
    return c


def main() -> None:
    pf, ph = paras(PATH_FINAL), paras(PATH_HUM)
    bf, bh = main_body_text(pf), main_body_text(ph)

    print("=== MAIN BODY (Ch 1-10) — phrase heuristics (NOT proof of authorship) ===\n")
    for label, text in ("FINAL", bf), ("HUMANIZED", bh):
        print(f"--- {label} ---")
        for pat, pname in (
            (FORBIDDEN, "forbidden/filler-ish"),
            (OVERCLAIM, "overclaim"),
            (STRUCTURAL_AI, "structural bridge phrases"),
        ):
            n, s = count_pattern(pat, text)
            print(f"  {pname}: {n} hits")
            for x in s[:4]:
                print(f"    … {x[:120]} …")
        print()

    print("=== US-spelling-shaped tokens (rough; many false positives) ===\n")
    for label, text in ("FINAL", bf), ("HUMANIZED", bh):
        tot = 0
        bits = []
        for pat, _ in US_VS_UK:
            n = len(re.findall(pat, text))
            if n:
                tot += n
                bits.append(f"{pat.strip(chr(92)+'b')}: {n}")
        print(f"{label}: ~{tot} token hits — {', '.join(bits[:6]) or 'none flagged'}")

    print("\n=== OOXML / layout (document.xml) ===\n")
    for label, pth in ("FINAL", PATH_FINAL), ("HUMANIZED", PATH_HUM):
        st = ooxml_stats(pth)
        print(label, st)

    print("\n=== Paragraph styles (first ~800 paras, python-docx) ===\n")
    for label, pth in ("FINAL", PATH_FINAL), ("HUMANIZED", PATH_HUM):
        c = style_names_sample(pth)
        print(label, c.most_common(12))

    print("\n=== Length ===\n")
    print("FINAL main body words:", len(bf.split()))
    print("HUM  main body words:", len(bh.split()))


if __name__ == "__main__":
    main()
