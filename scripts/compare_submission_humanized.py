"""Compare Final_Submission vs Humanized_version for coach QA (Chapters 1-10)."""
import collections
import re
import sys
from pathlib import Path

import docx

_ROOT = Path(__file__).resolve().parent.parent
PATH_ORIG = str(_ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission.docx")
PATH_HUM = str(
    _ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"
)


def paras_of(path: str) -> list[str]:
    return [p.text.strip() for p in docx.Document(path).paragraphs if p.text.strip()]


def split_chapters(paras: list[str]) -> collections.OrderedDict:
    hp = re.compile(r"^Chapter\s+(\d+)\s+[–-]")
    out: collections.OrderedDict = collections.OrderedDict()
    cur, buf = "front_matter", []
    for t in paras:
        m = hp.match(t)
        if m:
            if buf:
                out[cur] = buf
            cur = f"Chapter {m.group(1)}"
            buf = [t]
        else:
            buf.append(t)
    if buf:
        out[cur] = buf
    return out


def norm_num(s: str) -> str:
    return s.replace(",", "").replace(" ", "")


LOCKED_SUBSTRINGS = [
    "0.9989",
    "0.9984",
    "0.9986",
    "0.9996",
    "46.09",
    "0.9885",
    "0.9942",
    "22.70",
    "20.99",
    "920",
    "928",
    "934",
    "500,000",
    "2.3%",
    "97.7%",
    "128,002",
    "3.07",
    "31 MB",
    "187",
    "Dirichlet alpha = 0.5",
    "0.967",
    "0.983",
    "0.557",
    "0.990",
    "0.995",
    "0.973",
    "0.9923",
    "0.9961",
    "12.20",
    "17.42",
    "11.89",
    "15.60",
    "28.07",
    "0.484",
    "0.023",
    "0.0002",
    "0.0001",
]

FORBIDDEN = re.compile(
    r"delve|crucial|robust|comprehensive|leverage|utilize|facilitate|myriad|plethora|notably|underscores|pivotal|groundbreaking|seamlessly|nuanced understanding|it is important to note that|this highlights|in light of|overarching|paradigm",
    re.I,
)
OVERCLAIM = re.compile(
    r"\bprove[sd]?\b|\bproven\b|\bguaranteed\b|\bdefinitive\b|\bunquestionably\b|\beliminates\b|zero-risk|best overall",
    re.I,
)

PROTECTED = [
    "CICIoT2023",
    "kNN",
    "GAT",
    "GRU",
    "Random Forest",
    "MLP",
    "Flower",
    "FedAvg",
    "non-IID",
    "Dirichlet",
    "Captum",
    "Integrated Gradients",
    "FastAPI",
    "ECS-like",
    "ROC-AUC",
    "PyTorch Geometric",
    "Dynamic GNN",
    "SIEM",
    "SOC",
]

APPROVED_SNIPPETS = (
    "alabbadi",
    "albanbay",
    "basak",
    "cuppens",
    "miège",
    "miege",
    "han",
    "kokhlikyan",
    "kolias",
    "lazzarini",
    "lundberg",
    "lusa",
    "mcmahan",
    "ngo",
    "pinto",
    "qu",
    "sundararajan",
    "velickovic",
    "wang",
    "yang",
    "zheng",
    "zhong",
)


def cite_paren(text: str) -> list[str]:
    return re.findall(r"\([^)]*\d{4}[^)]*\)", text)


def cite_looks_external(cit: str) -> bool:
    low = cit.lower()
    if "chapter" in low or "section" in low:
        return False
    if "february" in low or "april" in low:
        return False
    return not any(a in low for a in APPROVED_SNIPPETS)


def abstract_block(paras: list[str]) -> tuple[str, int]:
    idx = next((i for i, t in enumerate(paras) if t.strip().lower() == "1. abstract"), None)
    if idx is None:
        return "", 0
    parts: list[str] = []
    for t in paras[idx + 1 :]:
        tl = t.lower()
        if tl.startswith("acknowledgements") or tl.startswith("chapter 1"):
            break
        parts.append(t)
    s = " ".join(parts)
    return s, len(s.split())


def table_text(doc) -> str:
    rows_out: list[str] = []
    for tbl in doc.tables:
        for row in tbl.rows:
            rows_out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(rows_out)


def main() -> int:
    po = paras_of(PATH_ORIG)
    ph = paras_of(PATH_HUM)
    co = split_chapters(po)
    ch = split_chapters(ph)

    ab_o, wc_o = abstract_block(po)
    ab_h, wc_h = abstract_block(ph)
    print("ABSTRACT_WORDS_ORIG", wc_o, "HUM", wc_h)

    issues: list[tuple[str, str]] = []
    flags: list[tuple[str, str, str]] = []

    for n in range(1, 11):
        k = f"Chapter {n}"
        if k not in co or k not in ch:
            issues.append((k, "MISSING_CHAPTER_IN_ONE_FILE"))
            continue
        to = "\n".join(co[k])
        th = "\n".join(ch[k])
        co_c = cite_paren(to)
        ch_c = cite_paren(th)
        if len(ch_c) < len(co_c):
            issues.append((k, f"CITATION_PAREN_DROP orig={len(co_c)} hum={len(ch_c)}"))

        for c in ch_c:
            if cite_looks_external(c):
                flags.append((k, "POSSIBLE_NON_APPROVED_CIT", c[:140]))

        to_norm = norm_num(to)
        th_norm = norm_num(th)
        for sub in LOCKED_SUBSTRINGS:
            sn = norm_num(sub)
            if sn in to_norm and sn not in th_norm:
                issues.append((k, f"LOCKED_SUBSTRING_MISSING: {sub}"))

        for m in FORBIDDEN.finditer(th):
            issues.append((k, f"FORBIDDEN_PHRASE: {m.group(0)}"))
        for m in OVERCLAIM.finditer(th):
            issues.append((k, f"OVERCLAIM_PHRASE: {m.group(0)}"))

        for term in PROTECTED:
            if term.lower() in to.lower() and term.lower() not in th.lower():
                issues.append((k, f"PROTECTED_TERM_LOST: {term}"))

    # Tables (numeric drift in Word tables)
    tbl_o = table_text(docx.Document(PATH_ORIG))
    tbl_h = table_text(docx.Document(PATH_HUM))
    for sub in ["0.9989", "0.9986", "22.70", "20.99", "0.967", "0.557"]:
        if sub in tbl_o and sub not in tbl_h:
            issues.append(("TABLES", f"CELL_VALUE_MISSING_IN_HUM: {sub}"))

    print("\n=== LIST 2 STYLE ISSUES (first 100) ===")
    for x in issues[:100]:
        print(x)
    print("TOTAL_ISSUES", len(issues))

    print("\n=== LIST 3 INTEGRITY FLAGS (first 60) ===")
    for x in flags[:60]:
        print(x)
    print("TOTAL_FLAGS", len(flags))

    # Chapter 10 voice
    k10 = "Chapter 10"
    if k10 in ch:
        t10 = "\n".join(ch[k10])
        i_n = len(re.findall(r"\bI\b", t10))
        print("\nCH10_I_COUNT", i_n)

    return 0


if __name__ == "__main__":
    sys.exit(main())
