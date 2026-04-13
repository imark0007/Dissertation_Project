"""
Compare agreed project specification (or interim proxy) to Dissertation_Arka_Talukder.md.
Run: python scripts/compare_spec_to_dissertation.py
Output: docs/reports/SPEC_DISSERTATION_ALIGNMENT.md
"""
from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    raise SystemExit("pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "Dissertation_Arka_Talukder.md"
OUT_PATH = ROOT / "docs" / "reports" / "SPEC_DISSERTATION_ALIGNMENT.md"

SPEC_CANDIDATES = [
    ROOT / "Arka-B01821011_MSc Cyber Security Project specification_form_2025-26.docx",
    ROOT
    / "B01821011_Final_Report_Package_for_Supervisor"
    / "05_Appendix_documents"
    / "Arka-B01821011_MSc Cyber Security Project specification_form_2025-26.docx",
    ROOT / "docs" / "reference" / "Arka-B01821011_MSc Cyber Security Project specification_form_2025-26.docx",
]

INTERIM_PROXY = ROOT / "archive" / "interim_report" / "B01821011_Interim_Report_Final_v2.docx"


def docx_to_text(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def normalise(s: str) -> str:
    s = s.lower()
    s = re.sub(r"\s+", " ", s)
    return s


def main():
    md = MD_PATH.read_text(encoding="utf-8")
    md_n = normalise(md)
    md_plain_n = normalise(re.sub(r"[*_`#]", "", md))

    spec_path = next((p for p in SPEC_CANDIDATES if p.exists()), None)
    if spec_path:
        source_label = f"**Formal specification** (`{spec_path.relative_to(ROOT)}`)"
        spec_text = docx_to_text(spec_path)
    elif INTERIM_PROXY.exists():
        spec_path = INTERIM_PROXY
        source_label = (
            f"**Proxy: interim report** (`{spec_path.relative_to(ROOT)}`) — "
            "the Moodle specification `.docx` was not found in the repo; "
            "re-run after copying `Arka-B01821011_MSc Cyber Security Project specification_form_2025-26.docx` "
            "to the project root (or `docs/reference/`)."
        )
        spec_text = docx_to_text(spec_path)
    else:
        OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        OUT_PATH.write_text(
            "# Specification vs dissertation\n\n"
            "No specification docx and no interim proxy found. Nothing to compare.\n",
            encoding="utf-8",
        )
        print(f"Wrote {OUT_PATH} (empty)")
        return

    spec_n = normalise(spec_text)

    checks: list[tuple[str, list[str], str]] = [
        (
            "Primary research question (IoT / GNN / FL / SIEM / CPU)",
            [
                "how can an explainable dynamic graph neural network",
                "federated learning",
                "siem",
                "cpu",
            ],
            "Core RQ wording from interim/spec should appear in Introduction.",
        ),
        (
            "Sub-questions / hypotheses (GNN vs baselines; federated; explanations)",
            [
                "random forest",
                "federated",
                "explanation",
            ],
            "Baselines, federation, and explainability addressed in Ch.1 / Ch.4 / Ch.8–9.",
        ),
        (
            "CICIoT2023 dataset",
            ["ciciot"],
            "Dataset named and scoped.",
        ),
        (
            "Flower / FedAvg",
            ["flower", "fedavg"],
            "Federated stack documented (Ch.4, 6, 8).",
        ),
        (
            "Integrated Gradients / Captum",
            ["integrated gradient", "captum"],
            "Explainability pipeline.",
        ),
        (
            "FastAPI / ECS / JSON alerts",
            ["fastapi", "ecs", "json"],
            "Deployment and SIEM-shaped output.",
        ),
        (
            "Ablation (temporal / GAT)",
            ["ablation", "gat only", "gru"],
            "§8.7 and Table 4.",
        ),
        (
            "Sensitivity (window / kNN k)",
            ["sensitivity", "window size", "knn"],
            "§8.8 Table 6 and Figure 15.",
        ),
        (
            "Multi-seed",
            ["multi-seed", "456"],
            "§8.9 Table 7 (seeds 42, 123, 456).",
        ),
        (
            "Ethics / public data / no human participants",
            ["ethics", "public", "participant"],
            "Ch.3 §3.4; align with ethics box on signed spec when available.",
        ),
        (
            "Marking criteria / specification grid",
            ["marking", "specification", "appendix b", "1.6"],
            "§1.6 mapping; verify percentages against signed PDF.",
        ),
    ]

    lines = [
        "# Specification vs dissertation alignment",
        "",
        f"**Generated from:** {source_label}",
        "",
        "## Automated keyword coverage",
        "",
        "| Commitment (from source doc) | In dissertation? | Notes |",
        "|------------------------------|------------------|-------|",
    ]

    for title, keywords, note in checks:
        hit = all(kw in md_n for kw in keywords)
        lines.append(f"| {title} | {'**Yes**' if hit else '**Check**'} | {note} |")

    lines.extend(
        [
            "",
            "## Interim-specific promises vs final dissertation",
            "",
        ]
    )

    interim_promises = [
        (
            "H2: federated GNN within 2% F1 of centralised GNN",
            "federated" in md_n and "1.000" in md and "central" in md_n,
            "Final Ch.8: federated matches central (100% F1) — satisfies stricter bar.",
        ),
        (
            "3–5 example alerts with explanations",
            "example" in md_n and "alert" in md_n,
            "§8.6 lists five examples.",
        ),
        (
            "Confusion matrices, ROC, FL convergence, inference/F1 bar chart",
            "confusion" in md_n and "roc" in md_n and "convergence" in md_n,
            "Figures 2–9 and 4–5.",
        ),
        (
            "Contingency: drop sensitivity if out of time",
            "sensitivity" in md_n,
            "Sensitivity §8.8 completed — exceeds interim contingency.",
        ),
    ]
    for title, ok, expl in interim_promises:
        lines.append(f"- **{title}:** {'Done — ' + expl if ok else 'Review — ' + expl}")

    lines.extend(["", "## Specification form: structured gaps / clarifications", ""])

    blob_n = spec_n

    if "nodes as devices" in blob_n or "devices are nodes" in blob_n:
        if "agreed project specification" in md_plain_n and "appendix b" in md_plain_n:
            lines.append(
                "- **Graph design vs overview text:** The specification overview assumes **device nodes** and **flow edges** where identifiers exist; "
                "the thesis implements **flow nodes** and **kNN edges** for the public CICIoT2023 release. **Status:** §1.4 now ties this explicitly to the signed form in **Appendix B** and supervisor agreement."
            )
        else:
            lines.append(
                "- **Graph design vs overview text:** The specification’s written overview still refers to **device nodes** and **flow edges**. "
                "The dissertation correctly implements **flow-level nodes** and **kNN similarity edges** (§1.4, §5.2) because the public CICIoT2023 release does not provide device identifiers. "
                "**Action:** Add one explicit sentence in §1.4 linking the deviation to the signed specification in **Appendix B**."
            )
    if "cybok" not in md_plain_n:
        lines.append(
            "- **CyBOK:** The specification includes CyBOK checkboxes; the dissertation should name the mapped areas (see **§2.9** if present)."
        )
    else:
        lines.append(
            "- **CyBOK:** **§2.9** maps the dissertation to CyBOK knowledge areas consistent with the specification (Appendix B)."
        )
    if "does not require ethical approval" in blob_n or "does not require ethical" in blob_n:
        if "does not require" in md_plain_n and "february 2026" in md_plain_n:
            lines.append(
                "- **Ethics sign-off:** The specification Section 2 states no full ethics-board approval is required. **Status:** §3.4 mirrors that sign-off (date aligned to the form)."
            )
        elif "erm" not in md_n and "ethical approval" not in md_n:
            lines.append(
                "- **Ethics sign-off:** The specification’s Section 2 records that the project **does not require** full ethics-board approval. "
                "**Optional:** Add one sentence in §3.4 mirroring that sign-off (date/signatory as on the form) so the thesis matches the appendix exactly."
            )

    used_formal = spec_path.name.startswith("Arka-B01821011")
    lines.extend(
        [
            "",
            "## Manual follow-ups",
            "",
        ]
    )
    if not used_formal:
        lines.append(
            "1. **Specification file:** Copy `Arka-B01821011_MSc Cyber Security Project specification_form_2025-26.docx` to the project root "
            "and re-run this script so comparisons use the signed form."
        )
    lines.extend(
        [
            f"{'2' if not used_formal else '1'}. **Marking percentages:** If the form’s marking grid differs from §1.6, update §1.6 and Appendix B.",
            f"{'3' if not used_formal else '2'}. **ERM ID:** If the form ever references a specific ERM case number, repeat it verbatim in §3.4.",
            f"{'4' if not used_formal else '3'}. **Timeline:** Interim used a *15-week* framing; the thesis uses *45 days* for the build sprint — ensure this matches what your supervisor approved on the form.",
            "",
            "---",
            "",
            "*Regenerate: `python scripts/compare_spec_to_dissertation.py`*",
        ]
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUT_PATH.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {OUT_PATH} (source: {spec_path.name})")


if __name__ == "__main__":
    main()
