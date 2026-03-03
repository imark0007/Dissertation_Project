"""
Convert Interim_Report_Arka_Talukder.md to Word (B01821011_Interim_Report_Final.docx).
- Replaces Figure 1 with figure1_pipeline.png
- Formats equations nicely (centered, equation style)
- Parses tables
Run from project root: python scripts/md_to_docx.py
Requires: pip install python-docx
"""
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Please install python-docx: pip install python-docx")
    raise

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "Interim_Report_Arka_Talukder.md"
OUT_PATH = PROJECT_ROOT / "B01821011_Interim_Report_Final.docx"
FIGURE_PATH = PROJECT_ROOT / "figure1_pipeline.png"


def set_cell_shading(cell, fill_color: str = "D9D9D9") -> None:
    """Set background shading for a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn('w:shd'):
            tcPr.remove(child)
            break
    tcPr.append(shading_elm)


def add_run_with_format(p, text):
    """Add text to paragraph, handling **bold** and *italic*."""
    remaining = text
    while remaining:
        b = remaining.find("**")
        i = remaining.find("*") if remaining.find("*") >= 0 else len(remaining)
        if b >= 0 and (b < i or i < 0):
            end = remaining.find("**", b + 2)
            if end >= 0:
                if b > 0:
                    p.add_run(remaining[:b])
                run = p.add_run(remaining[b + 2:end])
                run.bold = True
                remaining = remaining[end + 2:]
                continue
        if i >= 0 and not (i > 0 and remaining[i - 1] == "*"):
            end = remaining.find("*", i + 1)
            if end >= 0:
                if i > 0:
                    p.add_run(remaining[:i])
                run = p.add_run(remaining[i + 1:end])
                run.italic = True
                remaining = remaining[end + 1:]
                continue
        p.add_run(remaining)
        break


def add_equation(doc, eq_text: str) -> None:
    """Add a nicely formatted equation paragraph (centered, equation style)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(eq_text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.italic = True


def add_formatted_table(doc, headers: list, rows: list, caption: str = None) -> None:
    """Add a table with headers and rows."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(h)
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
        set_cell_shading(cell)
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data[:n_cols]):
            table.rows[i + 1].cells[j].text = str(val)
    if caption:
        p = doc.add_paragraph()
        p.add_run(caption).italic = True
        p.paragraph_format.space_before = Pt(6)


def parse_table(lines: list, i: int) -> tuple[list, list, int]:
    """Parse a markdown table. Returns (headers, rows, new_i)."""
    headers = [c.strip() for c in lines[i].split("|") if c.strip()]
    i += 1
    if i < len(lines) and re.match(r'^\|[\s\-:]+\|', lines[i]):
        i += 1  # skip separator
    rows = []
    while i < len(lines):
        row = lines[i]
        if not row.strip().startswith("|"):
            break
        parts = row.split("|")
        cells = [p.strip() for p in parts[1:-1]]  # drop leading/trailing empty from |
        if cells:
            rows.append(cells)
        i += 1
    return headers, rows, i


def latex_to_display(latex: str) -> str:
    """Convert simple LaTeX to display equation text."""
    s = latex.strip()
    s = re.sub(r'\\text\{([^}]+)\}', r'\1', s)
    s = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1) / (\2)', s)
    s = re.sub(r'F_1', 'F₁', s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip()


def parse_md_to_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name == "Normal":
            style.font.size = Pt(11)
            style.font.name = "Calibri"
            break

    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Explainable Dynamic Graph Neural Network SIEM for Software-Defined IoT using Edge AI and Federated Learning")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    doc.add_paragraph()
    for line in [
        "Author: Arka Talukder",
        "Student Number: B01821011",
        "Programme: MSc Cyber Security",
        "Supervisor: Dr. Raja Ujjan",
        "University of the West of Scotland",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    doc.add_page_break()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        # Markdown image: ![alt](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt, path = img_match.group(1), img_match.group(2)
            img_path = (md_path.parent / path).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.0))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(alt).italic = True
            i += 1
            continue

        # LaTeX equation block $$
        if stripped.startswith("$$") and stripped.endswith("$$"):
            eq_text = latex_to_display(stripped[2:-2])
            add_equation(doc, eq_text)
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped == "## References":
            doc.add_heading("References", level=1)
            i += 1
            while i < len(lines):
                ref_line = lines[i].strip()
                if ref_line:
                    doc.add_paragraph(ref_line)
                i += 1
            break

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and "|" in stripped[1:]:
            headers, rows, i = parse_table(lines, i)
            if headers and (rows or "---" in "".join(headers)):
                add_formatted_table(doc, headers, rows)
            continue

        if not stripped:
            i += 1
            continue

        # Skip code blocks (```)
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_run_with_format(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    parse_md_to_docx(MD_PATH, OUT_PATH)
