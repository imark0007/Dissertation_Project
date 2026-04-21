"""
Insert Final_Submission §1.6 and §2.2 into Humanized_version; renumber Humanized Ch 2 headings.

Save Humanized, then run:
  python scripts/align_humanized_to_final_structure.py
  python scripts/align_humanized_docx_errors.py
"""
from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
PATH_FINAL = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission.docx"
PATH_HUM = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"


def _clone_para_element(doc_final: Document, para_index: int):
    return deepcopy(doc_final.paragraphs[para_index]._element)


def _ch2_bounds(hum: Document) -> tuple[int, int]:
    ch2 = next(i for i, p in enumerate(hum.paragraphs) if p.text.strip().startswith("Chapter 2"))
    ch3 = next(i for i, p in enumerate(hum.paragraphs) if p.text.strip().startswith("Chapter 3"))
    return ch2, ch3


def patch_chapter1(hum: Document, final: Document) -> None:
    idx_16 = next(
        (i for i, p in enumerate(hum.paragraphs) if p.text.strip() == "1.6 Chapter Summary"),
        None,
    )
    if idx_16 is None:
        print("Skip Ch1: '1.6 Chapter Summary' not found (already aligned?)")
        return
    p0 = hum.paragraphs[idx_16]
    p1 = hum.paragraphs[idx_16 + 1]
    parent = p0._element.getparent()
    pos = parent.index(p0._element)
    parent.remove(p1._element)
    parent.remove(p0._element)
    for i, fi in enumerate(range(119, 124)):
        parent.insert(pos + i, _clone_para_element(final, fi))


def patch_chapter2_opening(hum: Document, final: Document) -> None:
    idx_21 = next(
        (i for i, p in enumerate(hum.paragraphs) if p.text.strip() == "2.1 Chapter Overview"),
        None,
    )
    if idx_21 is None:
        raise RuntimeError("2.1 Chapter Overview not found")
    p_body = hum.paragraphs[idx_21 + 1]
    if p_body.text.strip().startswith("2."):
        raise RuntimeError("Expected 2.1 body paragraph after overview heading")
    parent = p_body._element.getparent()
    pos = parent.index(p_body._element)
    parent.remove(p_body._element)
    parent.insert(pos, _clone_para_element(final, 126))
    parent.insert(pos + 1, _clone_para_element(final, 127))


def insert_themes_heading(hum: Document, final: Document) -> None:
    idx_iot = next(
        i
        for i, p in enumerate(hum.paragraphs)
        if p.text.strip().startswith("2.2 IoT Security and the Need for Detection")
    )
    iot_el = hum.paragraphs[idx_iot]._element
    parent = iot_el.getparent()
    pos = parent.index(iot_el)
    parent.insert(pos, _clone_para_element(final, 128))


def _ch2_already_aligned(hum: Document) -> bool:
    """True if §2.2 Themes exists and IoT is already §2.3 (re-run safe)."""
    seen_themes = False
    seen_iot_23 = False
    for p in hum.paragraphs:
        t = p.text.strip()
        if t == "2.2 Themes and structure":
            seen_themes = True
        if t.startswith("2.3 IoT Security and the Need for Detection"):
            seen_iot_23 = True
    return seen_themes and seen_iot_23


def bump_ch2_subsections(hum: Document) -> None:
    """Renumber 2.10→2.11 … 2.3→2.4 while IoT is still 2.2; then IoT 2.2→2.3."""
    ch2, ch3 = _ch2_bounds(hum)
    for n in range(10, 3, -1):
        prefix = f"2.{n} "
        for i in range(ch2, ch3):
            p = hum.paragraphs[i]
            raw = p.text
            st = raw.strip()
            if not st.startswith(prefix):
                continue
            # Before IoT is relabelled 2.2→2.3, only SIEM should be 2.3; never bump the IoT heading here.
            if n == 3 and "IoT Security and the Need for Detection" in st:
                continue
            p.text = f"2.{n + 1} " + st[len(prefix) :]
    for i, p in enumerate(hum.paragraphs):
        st = p.text.strip()
        if st.startswith("2.2 IoT Security and the Need for Detection"):
            p.text = "2.3 " + st[4:]
            break


def main() -> int:
    if not PATH_FINAL.is_file() or not PATH_HUM.is_file():
        print("Missing Final or Humanized .docx", file=sys.stderr)
        return 1

    final = Document(str(PATH_FINAL))
    hum = Document(str(PATH_HUM))

    patch_chapter1(hum, final)
    hum.save(str(PATH_HUM))

    hum = Document(str(PATH_HUM))
    final = Document(str(PATH_FINAL))
    if _ch2_already_aligned(hum):
        print("Skip Ch2: already aligned (2.2 Themes + 2.3 IoT present)")
    else:
        patch_chapter2_opening(hum, final)
        hum.save(str(PATH_HUM))

        hum = Document(str(PATH_HUM))
        final = Document(str(PATH_FINAL))
        insert_themes_heading(hum, final)
        hum.save(str(PATH_HUM))

        hum = Document(str(PATH_HUM))
        bump_ch2_subsections(hum)
        hum.save(str(PATH_HUM))

    print("OK:", PATH_HUM.name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
