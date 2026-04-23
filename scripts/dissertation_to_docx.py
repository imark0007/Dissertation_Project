"""
Convert Arka_Talukder_Dissertation_Final_DRAFT.md to Word for final submission.
Follows MSc Project Handbook: 1.5 line spacing, 11pt+ font, page numbers.
Embeds appendices (process docs, project spec) if available — see archive/README.md
for canonical paths to process + attendance under archive/process_attendance/.

After any edit to the thesis Markdown, prefer the bundled sync (also refreshes the
supervisor package copy of .md + .docx when that folder exists). By default this
runs the Final export and, if present, `Dissertation_Arka_Talukder_Humanized.md`
→ Humanized Word:

    python scripts/sync_dissertation_and_docx.py

Humanized-only:  python scripts/sync_humanized_md_and_docx.py

This file alone:      python scripts/dissertation_to_docx.py [--md PATH] [--out PATH]
Defaults: Arka_Talukder_Dissertation_Final_DRAFT.md → submission/Arka_Talukder_Dissertation_Final_DRAFT.docx
"""
import argparse
from pathlib import Path
import re
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION_START
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Install: pip install python-docx")
    raise

ROOT = Path(__file__).resolve().parent.parent
SUBMISSION_DIR = ROOT / "submission"
SUBMISSION_FORMS_DIR = SUBMISSION_DIR / "forms"
MD_PATH = ROOT / "Arka_Talukder_Dissertation_Final_DRAFT.md"
DEFAULT_OUT_PATH = SUBMISSION_DIR / "Arka_Talukder_Dissertation_Final_DRAFT.docx"

# Figure paths (relative to ROOT)
FIGURE_PATHS = {
    "assets/research_design_system.png": ROOT / "assets" / "research_design_system.png",
    "assets/figure1_pipeline.png": ROOT / "assets" / "figure1_pipeline.png",
    "assets/literature_positioning.png": ROOT / "assets" / "literature_positioning.png",
    "assets/literature_ids_taxonomy.png": ROOT / "assets" / "literature_ids_taxonomy.png",
    "assets/literature_dynamic_gnn_concept.png": ROOT / "assets" / "literature_dynamic_gnn_concept.png",
    "assets/literature_fedavg_flow.png": ROOT / "assets" / "literature_fedavg_flow.png",
    "assets/literature_explainability.png": ROOT / "assets" / "literature_explainability.png",
    "results/figures/similarity_knn_concept.png": ROOT / "results" / "figures" / "similarity_knn_concept.png",
    "results/figures/cm_gnn.png": ROOT / "results" / "figures" / "cm_gnn.png",
    "results/figures/roc_gnn.png": ROOT / "results" / "figures" / "roc_gnn.png",
    "results/figures/fl_convergence.png": ROOT / "results" / "figures" / "fl_convergence.png",
    "results/figures/model_comparison_bar.png": ROOT / "results" / "figures" / "model_comparison_bar.png",
    "results/figures/model_comparison.png": ROOT / "results" / "figures" / "model_comparison.png",
    "results/figures/sensitivity.png": ROOT / "results" / "figures" / "sensitivity.png",
    "results/figures/ablation_bar.png": ROOT / "results" / "figures" / "ablation_bar.png",
    "results/figures/cm_rf.png": ROOT / "results" / "figures" / "cm_rf.png",
    "results/figures/cm_mlp.png": ROOT / "results" / "figures" / "cm_mlp.png",
    "results/figures/roc_rf.png": ROOT / "results" / "figures" / "roc_rf.png",
    "results/figures/roc_mlp.png": ROOT / "results" / "figures" / "roc_mlp.png",
    # Appendix A — code screenshots (see scripts/render_appendix1_code_figures.py)
    "results/figures/appendix1/fig_a1_01_experiment_yaml.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_01_experiment_yaml.png",
    "results/figures/appendix1/fig_a1_02_preprocess.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_02_preprocess.png",
    "results/figures/appendix1/fig_a1_03_graph_knn.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_03_graph_knn.png",
    "results/figures/appendix1/fig_a1_04_graph_stratified.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_04_graph_stratified.png",
    "results/figures/appendix1/fig_a1_05_graph_sequence_dataset.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_05_graph_sequence_dataset.png",
    "results/figures/appendix1/fig_a1_06_dataloaders.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_06_dataloaders.png",
    "results/figures/appendix1/fig_a1_07_baselines.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_07_baselines.png",
    "results/figures/appendix1/fig_a1_08_dynamic_gnn.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_08_dynamic_gnn.png",
    "results/figures/appendix1/fig_a1_09_train_one_epoch.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_09_train_one_epoch.png",
    "results/figures/appendix1/fig_a1_10_explain_sequence.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_10_explain_sequence.png",
    "results/figures/appendix1/fig_a1_11_run_federated.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_11_run_federated.png",
    "results/figures/appendix1/fig_a1_12_flower_client.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_12_flower_client.png",
    "results/figures/appendix1/fig_a1_13_flower_server.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_13_flower_server.png",
    "results/figures/appendix1/fig_a1_14_fastapi_score.png": ROOT / "results" / "figures" / "appendix1" / "fig_a1_14_fastapi_score.png",
    # Chapter 6 — dark IDE-style code screenshots (see scripts/render_chapter6_code_screenshots.py)
    "results/figures/chapter6/fig_ch6_01_flows_to_knn_core.png": ROOT / "results" / "figures" / "chapter6" / "fig_ch6_01_flows_to_knn_core.png",
    "results/figures/chapter6/fig_ch6_02_stratified_split_core.png": ROOT / "results" / "figures" / "chapter6" / "fig_ch6_02_stratified_split_core.png",
    "results/figures/chapter6/fig_ch6_03_train_one_epoch.png": ROOT / "results" / "figures" / "chapter6" / "fig_ch6_03_train_one_epoch.png",
    "results/figures/chapter6/fig_ch6_04_dynamic_gnn_forward.png": ROOT / "results" / "figures" / "chapter6" / "fig_ch6_04_dynamic_gnn_forward.png",
    "results/figures/chapter6/fig_ch6_05_integrated_gradients_wrapper.png": ROOT / "results" / "figures" / "chapter6" / "fig_ch6_05_integrated_gradients_wrapper.png",
    "results/figures/chapter6/fig_ch6_06_fastapi_score_core.png": ROOT / "results" / "figures" / "chapter6" / "fig_ch6_06_fastapi_score_core.png",
}

# Appendix paths
PROCESS_DOC = ROOT / "archive" / "process_attendance" / "Arka_Talukder_Process_Documentation_B01821011.docx"
ATTENDANCE_DOC = ROOT / "archive" / "process_attendance" / "Arka Talukder_Attendance_Jan-Feb_B01821011.docx"
_SPEC_NAME = "Arka-B01821011_MSc Cyber Security Project specification_form_2025-26.docx"
_PROJECT_SPEC_CANDIDATES = [
    SUBMISSION_FORMS_DIR / _SPEC_NAME,
    ROOT / _SPEC_NAME,
    ROOT / "supervisor_package" / "05_Appendix_documents" / _SPEC_NAME,
    ROOT / "docs" / "reference" / _SPEC_NAME,
]


def _resolve_project_spec():
    for p in _PROJECT_SPEC_CANDIDATES:
        if p.exists():
            return p
    return _PROJECT_SPEC_CANDIDATES[0]


PROJECT_SPEC = _resolve_project_spec()

_FRONT_SHEET_CANDIDATES = [
    SUBMISSION_FORMS_DIR / "DissertationFrontSheet.docx",
    SUBMISSION_FORMS_DIR / "DissertationFrontSheet.doc",
    ROOT / "DissertationFrontSheet.docx",
    ROOT / "DissertationFrontSheet.doc",
    ROOT / "docs" / "reference" / "school_templates" / "DissertationFrontSheet.docx",
    ROOT / "docs" / "reference" / "school_templates" / "DissertationFrontSheet.doc",
]
_LIBRARY_FORM_CANDIDATES = [
    SUBMISSION_FORMS_DIR / "Dissertation Library Form.docx",
    SUBMISSION_FORMS_DIR / "Dissertation Library Form.DOC",
    ROOT / "Dissertation Library Form.docx",
    ROOT / "Dissertation Library Form.DOC",
    ROOT / "docs" / "reference" / "school_templates" / "Dissertation Library Form.docx",
    ROOT / "docs" / "reference" / "school_templates" / "Dissertation Library Form.DOC",
]
_DECLARATION_CANDIDATES = [
    ROOT / "Declaration of originality.docx",
    ROOT / "docs" / "reference" / "school_templates" / "Declaration of originality.docx",
]


def _resolve_first_existing(candidates: list[Path]) -> Path | None:
    for p in candidates:
        if p.exists():
            return p
    return None


def _add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_page_numbering(section, fmt: str = "decimal", start: int | None = None) -> None:
    """Set section page number format (decimal/roman) and optional restart."""
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def add_page_number_footer(section, *, fmt: str = "decimal"):
    """Add centered page number field with selected numbering format."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fmt == "roman":
        _add_field_run(p, "PAGE \\* roman")
    else:
        _add_field_run(p, "PAGE")


def add_front_pages(doc) -> None:
    """Create front sheet, declaration, and library release pages."""
    # Front page
    p = doc.add_paragraph("UWS LOGO (insert official logo here)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(11)
    doc.add_paragraph()
    title = doc.add_paragraph(
        "Explainable Dynamic Graph Neural Network SIEM for Software-Defined IoT using Edge AI and Federated Learning"
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
    doc.add_paragraph()
    for line in [
        "Arka Talukder",
        "Student ID: B01821011",
        "MSc Cyber Security (Full-time)",
        "School of Computing, Engineering and Physical Sciences",
        "University of the West of Scotland",
        "Supervisor: Dr. Raja Ujjan",
        "Moderator: Muhsin Hassanu",
        "Submission date: April 2026",
    ]:
        pr = doc.add_paragraph(line)
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Declaration page with red bordered block
    h = doc.add_paragraph("DECLARATION OF ORIGINALITY")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs:
        r.bold = True
        r.font.size = Pt(14)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell.text = (
        "I declare that this is an original study based on my own work and "
        "that I have not submitted it for any other course or degree."
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "12")
        elem.set(qn("w:color"), "C00000")
        borders.append(elem)
    tc_pr.append(borders)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].paragraph_format.space_before = Pt(10)
    cell.paragraphs[0].paragraph_format.space_after = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph("Signature: ________________________________")
    doc.add_paragraph("Name: Arka Talukder")
    doc.add_page_break()

    # Library release
    h2 = doc.add_paragraph("Library Release Form")
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h2.runs:
        r.bold = True
        r.font.size = Pt(14)
    doc.add_paragraph("Attach the signed UWS library release form on this page.")
    doc.add_paragraph("Reference: UWS MSc Project - Library Release Form")
    doc.add_page_break()


def add_auto_reference_pages(doc) -> None:
    """Insert Word auto fields for TOC, List of Figures, and List of Tables."""
    doc.add_paragraph("Table of Contents", style="Heading 1")
    p = doc.add_paragraph()
    _add_field_run(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    doc.add_paragraph("List of Figures", style="Heading 1")
    p2 = doc.add_paragraph()
    _add_field_run(p2, 'TOC \\h \\z \\c "Figure"')
    doc.add_page_break()

    doc.add_paragraph("List of Tables", style="Heading 1")
    p3 = doc.add_paragraph()
    _add_field_run(p3, 'TOC \\h \\z \\c "Table"')
    doc.add_page_break()


def add_placeholder_page(doc, title, instructions):
    """Add a placeholder page for Moodle forms."""
    doc.add_heading(title, level=1)
    doc.add_paragraph(instructions)
    doc.add_paragraph("[Insert downloaded form from Moodle here]")
    doc.add_page_break()


def resolve_image_path(md_path):
    """Resolve image path from markdown to filesystem."""
    for pattern, full_path in FIGURE_PATHS.items():
        if pattern in md_path and full_path.exists():
            return full_path
    # Fallback: try assets or results/figures
    if "assets/" in md_path:
        p = ROOT / md_path.replace("assets/", "assets/")
        if p.exists():
            return p
    if "results/figures/" in md_path:
        name = md_path.split("/")[-1]
        p = ROOT / "results" / "figures" / name
        if p.exists():
            return p
    return None


def _markdown_plain(s: str) -> str:
    """Strip inline markdown / link syntax for clean submission-style Word text."""
    if not s:
        return ""
    t = s.replace("\u00a0", " ").replace("\u200b", "")
    t = t.replace("—", " - ")
    t = t.replace("§", "Section ")
    # [label](url) -> label (no raw brackets in Word/PDF body)
    t = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", t)
    while "**" in t:
        n = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
        if n == t:
            break
        t = n
    while "`" in t:
        n = re.sub(r"`([^`]+)`", r"\1", t)
        if n == t:
            break
        t = n
    while "*" in t:
        n = re.sub(r"\*([^*]+)\*", r"\1", t)
        if n == t:
            break
        t = n
    return t.strip()


def _apply_body_text_format(p) -> None:
    """Enforce 11pt Times New Roman on paragraph runs (lists inherit Normal inconsistently)."""
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    if not p.runs and p.text:
        # Edge case: single run without .runs populated yet
        pass


def _add_paragraph(
    doc,
    text: str,
    style: str | None = None,
    *,
    italic: bool = False,
):
    """Add a body paragraph; optional built-in style and whole-paragraph italic."""
    t = _markdown_plain(text)
    if not t:
        return None
    st = style if style and style in doc.styles else None
    p = doc.add_paragraph(t, style=st) if st else doc.add_paragraph(t)
    if italic:
        for r in p.runs:
            r.italic = True
    _apply_body_text_format(p)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def _list_bullet_style(level: int) -> str:
    names = ["List Bullet", "List Bullet 2", "List Bullet 3"]
    return names[min(max(level, 0), len(names) - 1)]


def _set_table_borders(table, *, color: str = "000000", size: str = "8") -> None:
    """Apply visible single-line borders to all table edges and inner grid."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.append(tbl_pr)

    # Remove existing borders so regenerated documents stay consistent.
    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)
        borders.append(elem)
    tbl_pr.append(borders)


def _add_bordered_picture(doc, image_path: Path, *, width: Inches = Inches(5.5)) -> None:
    """
    Insert image inside a 1x1 table cell so Word always shows a clean figure border.
    This mirrors the boxed figure look in sample-style dissertations.
    """
    fig_tbl = doc.add_table(rows=1, cols=1)
    fig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if "Table Grid" in doc.styles:
        fig_tbl.style = "Table Grid"
    _set_table_borders(fig_tbl, color="000000", size="8")
    cell = fig_tbl.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=width)


def process_markdown(doc, content):
    """Parse markdown and add to document."""
    lines = content.split("\n")
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Skip front matter placeholder
        if "## Front Matter" in line:
            i += 1
            while i < len(lines) and not lines[i].startswith("## ") and not (lines[i].startswith("# ") and "Abstract" not in lines[i]):
                i += 1
            if i < len(lines) and "Table of Contents" in lines[i]:
                i += 1
            continue
        # Skip markdown TOC / list sections (Word fields are inserted separately)
        if line.strip() == "## Table of Contents":
            i += 1
            while i < len(lines) and not lines[i].startswith("## Chapter 1"):
                i += 1
            continue

        # Headers (strip markdown so Word shows no stray ** or `)
        if line.startswith("# "):
            h = doc.add_heading(_markdown_plain(line[2:]), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in h.runs:
                r.font.size = Pt(16)
                r.bold = True
        elif line.startswith("## "):
            htxt = _markdown_plain(line[3:])
            h = doc.add_heading(htxt, level=1)
            if htxt.startswith("Chapter "):
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in h.runs:
                    r.font.size = Pt(16)
                    r.bold = True
                    r.italic = False
        elif line.startswith("### "):
            doc.add_heading(_markdown_plain(line[4:]), level=2)
        elif line.startswith("#### "):
            # Word Heading 3 (## = H1, ### = H2, #### = H3)
            doc.add_heading(_markdown_plain(line[5:]), level=3)
        elif line.startswith("##### "):
            doc.add_heading(_markdown_plain(line[6:]), level=4)

        elif line.strip() == "---":
            pass

        # Table
        elif "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        else:
            if in_table and table_lines:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                rows = []
                for tl in table_lines[2:]:
                    if "|" in tl:
                        cells = [c.strip() for c in tl.split("|")[1:-1]]
                        rows.append(cells)
                if headers and rows:
                    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    if "Table Grid" in doc.styles:
                        t.style = "Table Grid"
                    _set_table_borders(t, color="000000", size="8")
                    for j, h in enumerate(headers):
                        t.rows[0].cells[j].text = _markdown_plain(h)
                    # Estimated page numbers for Table of Figures and Tables
                    fig_pages = {
                        1: 28,
                        2: 42,
                        3: 43,
                        4: 45,
                        5: 46,
                        6: 47,
                        7: 48,
                        8: 49,
                        9: 50,
                        10: 22,
                        11: 23,
                        12: 24,
                        13: 25,
                        14: 26,
                        15: 52,
                        16: 53,
                        17: 31,
                        18: 55,
                        19: 56,
                        20: 57,
                        21: 58,
                        22: 59,
                        23: 60,
                    }
                    table_pages = {1: 42, 2: 45, 3: 46}
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            if ci < len(t.rows[ri + 1].cells):
                                val = cell
                                if (val.strip() in ("—", "–", "-") or val == "—") and ci == len(headers) - 1 and "Page" in str(headers):
                                    if "Figure" in str(headers):
                                        val = str(fig_pages.get(ri + 1, "—"))
                                    elif "Table" in str(headers):
                                        val = str(table_pages.get(ri + 1, "—"))
                                t.rows[ri + 1].cells[ci].text = _markdown_plain(val)
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                for r in p.runs:
                                    r.font.name = "Times New Roman"
                                    r.font.size = Pt(11)
                in_table = False
                table_lines = []

            # Image (avoid duplicate captions when MD uses ![...] then **Figure N: ...**)
            if "![" in line and "]" in line and "(" in line:
                match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if match:
                    img_path = resolve_image_path(match.group(2))
                    if img_path and img_path.exists():
                        _add_bordered_picture(doc, img_path, width=Inches(5.5))
                        alt = match.group(1).strip()
                        j = i + 1
                        while j < len(lines) and not lines[j].strip():
                            j += 1
                        cap_line = lines[j].strip() if j < len(lines) else ""
                        if cap_line and (
                            re.match(r"^\*\*Figure \d+:", cap_line)
                            or re.match(r"^\*\*Figure A1-\d+", cap_line)
                        ):
                            doc.add_paragraph(_markdown_plain(lines[j]), style="Caption")
                            i = j  # skip blank lines + caption line; outer loop also increments i
                        elif alt:
                            doc.add_paragraph(_markdown_plain(alt), style="Caption")
                    else:
                        doc.add_paragraph(
                            _markdown_plain(f"[Figure: {match.group(1)} — image not found at {match.group(2)}]")
                        )
            elif (mb := re.match(r"^(\s*)-\s+(.+)$", line)):
                lead, body = mb.group(1), mb.group(2)
                lvl = min(len(lead.replace("\t", "    ")) // 2, 2)
                lb = _list_bullet_style(lvl)
                _add_paragraph(doc, body, style=lb if lb in doc.styles else None)
            elif (mn := re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)):
                body = mn.group(3)
                lvl = min(len(mn.group(1).replace("\t", "    ")) // 2, 2)
                ln = ["List Number", "List Number 2", "List Number 3"][lvl]
                _add_paragraph(doc, body, style=ln if ln in doc.styles else None)
            elif line.strip().startswith("*Source") or (
                line.strip().startswith("*") and "Source:" in line
            ):
                _add_paragraph(doc, line, italic=True)
            elif line.strip() and not line.strip().startswith("*Note:"):
                _add_paragraph(doc, line)

        i += 1


def append_document(master_doc, other_path):
    """Append another docx to master. Uses docxcompose if available."""
    try:
        from docxcompose.composer import Composer
        composer = Composer(master_doc)
        other = Document(str(other_path))
        composer.append(other)
        return True
    except ImportError:
        try:
            other = Document(str(other_path))
            for element in other.element.body:
                master_doc.element.body.append(element)
            return True
        except Exception as e:
            print(f"Warning: Could not append {other_path}: {e}")
            return False
    except Exception as e:
        print(f"Warning: Could not append {other_path}: {e}")
        return False


def _convert_legacy_doc_to_docx(src: Path, dst: Path) -> bool:
    """
    Convert legacy .doc/.DOC to .docx using local Word COM automation (Windows).
    Returns True on success.
    """
    dst.parent.mkdir(parents=True, exist_ok=True)
    # SaveAs2 FileFormat=16 => wdFormatDocumentDefault (.docx)
    src_ps = str(src.resolve()).replace("'", "''")
    dst_ps = str(dst.resolve()).replace("'", "''")
    ps = (
        "$ErrorActionPreference='Stop';"
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        "$word.DisplayAlerts=0;"
        f"$doc=$word.Documents.Open('{src_ps}');"
        f"$doc.SaveAs2('{dst_ps}',16);"
        "$doc.Close();"
        "$word.Quit();"
    )
    try:
        run = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", ps],
            check=False,
            capture_output=True,
            text=True,
        )
    except Exception as e:
        print(f"Warning: could not run Word conversion for {src}: {e}")
        return False
    if run.returncode != 0:
        stderr = (run.stderr or "").strip()
        if stderr:
            print(f"Warning: Word conversion failed for {src}: {stderr[:240]}")
        else:
            print(f"Warning: Word conversion failed for {src} (exit {run.returncode})")
        return False
    return dst.exists()


def _append_or_note_form(doc, title: str, source_path: Path | None, generated_dir: Path) -> None:
    """Append form document when possible; otherwise keep explicit placeholder note."""
    # Keep explicit section title so QA checks and human readers can verify form presence.
    doc.add_heading(title, level=1)
    if source_path is None:
        doc.add_paragraph(f"Required form not found automatically. Insert manually: {title}")
        doc.add_page_break()
        return

    path_for_append = source_path
    if source_path.suffix.lower() == ".doc":
        converted = generated_dir / f"{source_path.stem}.docx"
        if _convert_legacy_doc_to_docx(source_path, converted):
            path_for_append = converted
        else:
            doc.add_paragraph(f"Could not auto-convert legacy file. Insert manually from: {source_path}")
            doc.add_page_break()
            return

    # If .docx path is mergeable, append directly.
    if path_for_append.suffix.lower() == ".docx":
        ok = append_document(doc, path_for_append)
        if not ok:
            doc.add_paragraph(f"Could not merge form automatically. Insert manually from: {path_for_append}")
            doc.add_page_break()
        else:
            doc.add_page_break()
    else:
        doc.add_paragraph(f"Unsupported form type. Insert manually from: {source_path}")
        doc.add_page_break()


def add_official_front_forms(doc) -> bool:
    """
    Insert official school front forms when available.
    Returns True if at least one form was inserted or placeholder-added via this routine.
    """
    front_sheet = _resolve_first_existing(_FRONT_SHEET_CANDIDATES)
    declaration = _resolve_first_existing(_DECLARATION_CANDIDATES)
    library_form = _resolve_first_existing(_LIBRARY_FORM_CANDIDATES)
    if not any([front_sheet, declaration, library_form]):
        return False

    generated_dir = ROOT / "results" / "generated_forms"
    _append_or_note_form(doc, "Dissertation Front Sheet", front_sheet, generated_dir)
    _append_or_note_form(doc, "Declaration of Originality", declaration, generated_dir)
    _append_or_note_form(doc, "Library Release Form", library_form, generated_dir)
    return True


def main(out_path: Path, md_path: Path | None = None) -> None:
    src = md_path if md_path is not None else MD_PATH
    content = src.read_text(encoding="utf-8")
    doc = Document()

    # Set default font and line spacing (UWS MSc final report: 11pt body, 1.5 line spacing)
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(11)
    font.name = "Times New Roman"
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading styles (submission-oriented defaults).
    if "Heading 1" in doc.styles:
        doc.styles["Heading 1"].font.size = Pt(14)
        doc.styles["Heading 1"].font.name = "Times New Roman"
        doc.styles["Heading 1"].font.bold = True
        doc.styles["Heading 1"].font.italic = False
    if "Heading 2" in doc.styles:
        doc.styles["Heading 2"].font.size = Pt(12)
        doc.styles["Heading 2"].font.name = "Times New Roman"
        doc.styles["Heading 2"].font.bold = True
        doc.styles["Heading 2"].font.italic = False
    if "Heading 3" in doc.styles:
        doc.styles["Heading 3"].font.size = Pt(12)
        doc.styles["Heading 3"].font.name = "Times New Roman"
        doc.styles["Heading 3"].font.bold = False
        doc.styles["Heading 3"].font.italic = True
    if "Heading 4" in doc.styles:
        doc.styles["Heading 4"].font.size = Pt(11)
        doc.styles["Heading 4"].font.name = "Times New Roman"
        doc.styles["Heading 4"].font.bold = False
        doc.styles["Heading 4"].font.italic = True

    # Heading spacing (readable like programme sample reports; tweak in Word if needed)
    _h_space = [
        ("Heading 1", Pt(14), Pt(10)),
        ("Heading 2", Pt(12), Pt(8)),
        ("Heading 3", Pt(10), Pt(6)),
        ("Heading 4", Pt(8), Pt(4)),
    ]
    for hn, before, after in _h_space:
        if hn in doc.styles:
            pf = doc.styles[hn].paragraph_format
            pf.space_before = before
            pf.space_after = after
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if "Caption" in doc.styles:
        cs = doc.styles["Caption"]
        cs.font.name = "Times New Roman"
        cs.font.size = Pt(10)
        cs.font.italic = True
        cs.paragraph_format.space_after = Pt(6)
        cs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Page margins (left bound edge wider for print binding).
    for section in doc.sections:
        section.top_margin = Inches(1.0)      # 2.54 cm
        section.bottom_margin = Inches(1.0)   # 2.54 cm
        section.right_margin = Inches(1.0)    # 2.54 cm
        section.left_margin = Inches(1.25)    # 3.17 cm

    # Front-matter pages aligned to submission checklist.
    # Prefer official school forms when available; otherwise use generated placeholders.
    used_official_forms = add_official_front_forms(doc)
    if not used_official_forms:
        add_front_pages(doc)
    add_auto_reference_pages(doc)

    # Front matter numbering (roman)
    set_page_numbering(doc.sections[0], fmt="roman", start=1)
    add_page_number_footer(doc.sections[0], fmt="roman")

    # Main content starts in a new section (arabic numbering)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_numbering(doc.sections[-1], fmt="decimal", start=1)
    add_page_number_footer(doc.sections[-1], fmt="decimal")
    process_markdown(doc, content)

    # Embedded .docx submissions (must NOT reuse "Appendix A/B/C" here — those letters
    # are already used in the Markdown body for manifest + reproducibility + code figures.)
    if PROCESS_DOC.exists():
        doc.add_page_break()
        append_document(doc, PROCESS_DOC)

    if ATTENDANCE_DOC.exists():
        doc.add_page_break()
        append_document(doc, ATTENDANCE_DOC)

    if PROJECT_SPEC.exists():
        doc.add_page_break()
        append_document(doc, PROJECT_SPEC)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_path)
    except PermissionError:
        pkg_dir = ROOT / "supervisor_package" / "01_Dissertation"
        pkg = pkg_dir / out_path.name
        if out_path.resolve() != pkg.resolve() and pkg_dir.is_dir():
            print(f"Permission denied writing {out_path}; trying {pkg}", file=sys.stderr)
            doc.save(pkg)
            print(f"Saved: {pkg}")
        else:
            raise
    else:
        print(f"Saved: {out_path}")
    if not PROJECT_SPEC.exists():
        print(
            "WARNING: Agreed specification .docx not found. Embed manually from:",
            [str(p.relative_to(ROOT)) for p in _PROJECT_SPEC_CANDIDATES],
        )
    print("Next steps:")
    if used_official_forms:
        print("  1. Verify front sheet, declaration, and library forms render correctly")
    else:
        print("  1. Replace placeholder pages with downloaded forms from Moodle")
    print("  2. Update TOC / List of Figures / List of Tables page numbers in Word (Insert > Table of Contents)")
    print("  3. Abstract: add drop cap on first letter if required by the School sample")
    print("  4. List of Abbreviations: colour the abbreviation column in Word if required")
    print("  5. Body text is 11pt / 1.5 spacing from Normal style; verify in Word; submit per module")
    print("  6. Final Harvard list: Mendeley / Zotero / EndNote; see docs/reports/FINAL_WORD_FORMAT_MENDELEY_PAGE_BUDGET.md")


def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Convert dissertation Markdown to Word.")
    p.add_argument(
        "--md",
        type=Path,
        default=None,
        help=f"Input .md path (default: {MD_PATH.name})",
    )
    p.add_argument(
        "--out",
        type=Path,
        default=None,
        help=f"Output .docx path (default: {DEFAULT_OUT_PATH.name} in repo root)",
    )
    return p.parse_args()


if __name__ == "__main__":
    args = _parse_args()
    out = (args.out if args.out is not None else DEFAULT_OUT_PATH).resolve()
    md = (args.md.resolve() if args.md is not None else MD_PATH)
    main(out_path=out, md_path=md)
