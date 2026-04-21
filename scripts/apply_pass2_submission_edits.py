"""Pass 2 surgical edits: Arka_Talukder_Dissertation_Final_Submission.docx"""
from pathlib import Path

import docx

PATH = Path(__file__).resolve().parents[1] / "submission" / "Arka_Talukder_Dissertation_Final_Submission.docx"

TITLE = (
    "Explainable Dynamic Graph Neural Network SIEM for Software-Defined IoT "
    "using Edge AI and Federated Learning"
)


def main() -> None:
    d = docx.Document(str(PATH))

    # --- Front sheet placeholders (paragraph indices from Pass 1 audit) ---
    d.paragraphs[3].text = TITLE
    d.paragraphs[6].text = "Not applicable"
    d.paragraphs[10].text = "Arka Talukder"
    d.paragraphs[11].text = "B01821011"
    d.paragraphs[19].text = "in Cyber Security"
    d.paragraphs[24].text = "23 April 2026"

    # --- Chapter 8.2: align figure cross-references with formal captions ---
    p360 = d.paragraphs[360].text
    old_a = (
        "Figures 17 and 18 show confusion matrices for Random Forest and MLP. "
        "Figures 19 and 20 show ROC curves for RF and MLP."
    )
    new_a = (
        "Figures 18 and 19 show confusion matrices for Random Forest and MLP. "
        "Figures 20 and 21 show ROC curves for RF and MLP."
    )
    if old_a not in p360:
        raise SystemExit("Expected sentence block not found in paragraph 360; abort.")
    d.paragraphs[360].text = p360.replace(old_a, new_a)

    p375 = d.paragraphs[375].text
    old_b = (
        "Figures 17 and 18 show 187 FP for Random Forest and 4 FP for MLP "
        "(benign predicted as attack)."
    )
    new_b = (
        "Figures 18 and 19 show 187 FP for Random Forest and 4 FP for MLP "
        "(benign predicted as attack)."
    )
    if old_b not in p375:
        raise SystemExit("Expected sentence not found in paragraph 375; abort.")
    d.paragraphs[375].text = p375.replace(old_b, new_b)

    d.save(str(PATH))
    print("Saved:", PATH)


if __name__ == "__main__":
    main()
