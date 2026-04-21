"""Count words per chapter from dissertation DOCX."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "submission" / "Arka_Talukder_Dissertation_Final.docx"


def count_words(s: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+", s))


def main() -> None:
    doc = Document(str(DOCX))
    paras = [p.text.strip() for p in doc.paragraphs]
    heads: list[tuple[int, int, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        style_name = p.style.name if p.style else ""
        if not style_name.startswith("Heading"):
            continue
        m = re.match(r"^Chapter\s+(\d+)\s*[–-]\s*(.+)$", txt, flags=re.I)
        if m:
            heads.append((int(m.group(1)), idx, txt))
    # Keep first heading occurrence per chapter number.
    dedup: dict[int, tuple[int, int, str]] = {}
    for row in sorted(heads, key=lambda x: x[1]):
        dedup.setdefault(row[0], row)
    heads = sorted(dedup.values(), key=lambda x: x[0])
    print(DOCX.name)
    for i, (ch, start, title) in enumerate(heads):
        end = heads[i + 1][1] if i + 1 < len(heads) else len(paras)
        block = " ".join(paras[start:end])
        print(f"  {title}: {count_words(block)}")


if __name__ == "__main__":
    main()
