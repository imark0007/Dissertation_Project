"""
Approved batch: Pass 4 note before Table 2, §2.2 insert (main submission only),
citation standardisation (skip Chapter 10), school line on front sheet.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
PATH_SUB = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission.docx"
PATH_HUM = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"

PASS4_NOTE = (
    "A note on interpretation. The F1 and ROC-AUC values reported below reach a ceiling "
    "of 1.0 on this test split. This is reported as observed evidence, not as a claim of "
    "universal performance. A perfect score on a 934-sequence test set drawn from a single "
    "CICIoT2023 subset is consistent with two non-mutually-exclusive readings: (a) the binary "
    "benign-vs-attack task is well-separated by the engineered flow features at this scale, "
    "and (b) the stratified windowing produces sequences whose label is highly predictable once "
    "the GAT has converged. Section 10.5 returns to this honestly, and Sections 8.7–8.9 "
    "(ablation, sensitivity, multi-seed) are included precisely so the ceiling is not a "
    "single-configuration artefact."
)

SECTION_22_HEADING = "2.2 Themes and structure"
SECTION_22_BODY = (
    "The literature is grouped by theme so each block ties to the research questions in "
    "Section 1.3. Section 2.3 covers IoT threat context and CICIoT2023; Section 2.4 covers "
    "SIEM/SOC workflows and alert quality; Sections 2.5 and 2.6 cover graph and dynamic GNN "
    "ideas and federated learning; Section 2.7 covers explainability in ML-based security; "
    "Sections 2.8 to 2.10 cover the gap, CyBOK mapping, extended comparison, and summary. "
    "Technical numbers and experiment design are left to later chapters so this review does "
    "not duplicate methodology."
)


def insert_paragraph_before(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc: Document, predicate) -> Paragraph | None:
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    return None


def insert_pass4_before_table2(doc: Document) -> bool:
    if any("A note on interpretation" in p.text for p in doc.paragraphs):
        return False
    target = find_para(doc, lambda t: t.startswith("Table 2:") and "Model comparison" in t)
    if not target:
        print("ERROR: Table 2 caption not found", file=sys.stderr)
        return False
    insert_paragraph_before(target, PASS4_NOTE)
    return True


def insert_section_22_before_23(doc: Document) -> bool:
    if any(p.text.strip().startswith("2.2 Themes and structure") for p in doc.paragraphs):
        return False
    target = find_para(doc, lambda t: t.startswith("2.3 IoT Security and the Need for Detection"))
    if not target:
        print("WARN: 2.3 IoT heading not found; skip §2.2 insert", file=sys.stderr)
        return False
    insert_paragraph_before(target, SECTION_22_BODY)
    insert_paragraph_before(target, SECTION_22_HEADING)
    return True


def standardise_citations(doc: Document) -> int:
    ch = None
    n = 0
    for p in doc.paragraphs:
        t = p.text
        if t.startswith("Chapter "):
            m = re.match(r"^Chapter\s+(\d+)", t)
            ch = int(m.group(1)) if m else None
        if ch == 10:
            continue
        orig = t
        t = re.sub(r"Alabbadi and Alabbadi and Bajaber \(2025\)", "Alabbadi and Bajaber (2025)", t)
        t = re.sub(r"(?<!Alabbadi and )Bajaber \(2025\)", "Alabbadi and Bajaber (2025)", t)
        t = re.sub(r"(?<!Lundberg and )Lee \(2017\)", "Lundberg and Lee (2017)", t)
        if t != orig:
            p.text = t
            n += 1
    return n


def fix_front_sheet_school(doc: Document) -> bool:
    for p in doc.paragraphs:
        if p.text.strip() == "School of Computing":
            p.text = "School of Computing, Engineering and Physical Sciences"
            return True
    return False


def process_submission(path: Path) -> None:
    doc = Document(str(path))
    r4 = insert_pass4_before_table2(doc)
    r2 = insert_section_22_before_23(doc)
    n = standardise_citations(doc)
    fs = fix_front_sheet_school(doc)
    doc.save(str(path))
    print(path.name, "Pass4:", r4, "§2.2:", r2, "citation_paras:", n, "school_line:", fs)


def process_humanized(path: Path) -> None:
    doc = Document(str(path))
    r4 = insert_pass4_before_table2(doc)
    n = standardise_citations(doc)
    fs = fix_front_sheet_school(doc)
    doc.save(str(path))
    print(path.name, "Pass4:", r4, "§2.2: skipped (already numbered)", "citation_paras:", n, "school_line:", fs)


def main() -> int:
    if not PATH_SUB.is_file():
        print("Missing", PATH_SUB, file=sys.stderr)
        return 1
    process_submission(PATH_SUB)
    if PATH_HUM.is_file():
        process_humanized(PATH_HUM)
    else:
        print("Skip humanized (file not found)", PATH_HUM)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
