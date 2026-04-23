"""
Rebuild Dissertation_Arka_Talukder_Humanized.md from the Humanized Word export.

- Pulls Acknowledgements, List of Abbreviations, Abstract, and all chapters from
  Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx.
- Preserves the manual Table of Contents / List of Figures / List of Tables block from the
  existing Humanized .md (Word stores TOC fields without plain-text entries).
- Maps embedded images to repo-relative paths via SHA256 where possible; falls back to the
  figure order in the current Humanized .md for unmapped blips.

Usage:

    python scripts/export_humanized_docx_to_md.py
    python scripts/export_humanized_docx_to_md.py --dry-run
"""
from __future__ import annotations

import argparse
import hashlib
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOCX = ROOT / "submission" / "Arka_Talukder_Dissertation_Final_Submission_Humanized_version.docx"
DEFAULT_MD = ROOT / "Dissertation_Arka_Talukder_Humanized.md"
# Overridden by CLI in main()
DOCX_PATH = DEFAULT_DOCX
MD_PATH = DEFAULT_MD


def iter_block_items(parent: DocumentObject):
    for el in parent.element.body:
        if el.tag == qn("w:p"):
            yield Paragraph(el, parent)
        elif el.tag == qn("w:tbl"):
            yield Table(el, parent)


def norm_text(s: str) -> str:
    return (
        s.replace("\u2013", "–")
        .replace("\u2014", "—")
        .replace("\u00a0", " ")
        .replace("\ufffd", "–")
    )


def paragraph_to_md(p: Paragraph) -> str:
    parts: list[str] = []
    for r in p.runs:
        t = norm_text(r.text or "")
        if not t:
            continue
        if r.bold:
            parts.append(f"**{t}**")
        elif r.italic:
            parts.append(f"*{t}*")
        else:
            parts.append(t)
    return "".join(parts).strip()


def table_is_figure_frame(table: Table) -> bool:
    if len(table.rows) != 1 or len(table.rows[0].cells) != 1:
        return False
    cell = table.rows[0].cells[0]
    for p in cell.paragraphs:
        for r in p.runs:
            el = r._element
            if el.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
            ):
                return True
            if el.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor"
            ):
                return True
    return False


def _build_repo_hash_index() -> dict[str, Path]:
    out: dict[str, Path] = {}
    for p in list(ROOT.glob("assets/*.png")) + list(ROOT.glob("results/figures/**/*.png")):
        h = hashlib.sha256(p.read_bytes()).hexdigest()
        out[h] = p
    return out


def _resolve_blip_paths(docx_path: Path, by_hash: dict[str, Path]) -> list[str | None]:
    z = zipfile.ZipFile(docx_path)
    rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
    id_to_target = dict(re.findall(r'Id="(rId\d+)"[^>]*Target="([^"]+)"', rels))
    xml = z.read("word/document.xml").decode("utf-8")
    rids = re.findall(r'<a:blip[^>]*r:embed="(rId\d+)"', xml)
    out: list[str | None] = []
    for rid in rids:
        tgt = id_to_target.get(rid, "")
        if not tgt:
            out.append(None)
            continue
        tgt = tgt.replace("../", "")
        if not tgt.startswith("word/"):
            tgt = "word/" + tgt
        if tgt not in z.namelist():
            out.append(None)
            continue
        h = hashlib.sha256(z.read(tgt)).hexdigest()
        p = by_hash.get(h)
        if p is None:
            out.append(None)
        else:
            out.append(str(p.relative_to(ROOT)).replace("\\", "/"))
    return out


def _figure_fallbacks_from_md(md_text: str) -> list[str]:
    return [m.group(1).strip() for m in re.finditer(r"!\[[^\]]*\]\(([^)]+)\)", md_text)]


def _strip_trailing_horizontal_rules(md: str) -> str:
    s = md.rstrip().replace("\r\n", "\n")
    while True:
        ns = re.sub(r"(\n---\s*)+$", "\n", s).rstrip()
        if ns == s:
            break
        s = ns
    return s + "\n" if s else ""


def _extract_preserved_toc_block(md_text: str) -> str:
    start = md_text.find("\n## Table of Contents\n")
    if start == -1:
        return ""
    start += 1
    end = md_text.find("\n## Chapter 1", start)
    if end == -1:
        return ""
    return md_text[start:end].rstrip() + "\n"


def _extract_abbrev_note(md_text: str) -> str:
    """Italic note line after abbrev table in existing MD, if any."""
    m = re.search(
        r"(## List of Abbreviations\n.*?\n\n)(\*Note:[^\n]+\n)",
        md_text,
        re.S,
    )
    return m.group(2).rstrip() if m else ""


def _title_from_docx(doc: DocumentObject) -> str:
    title = (
        "Explainable Dynamic Graph Neural Network SIEM for Software-Defined IoT "
        "using Edge AI and Federated Learning"
    )
    for p in doc.paragraphs[:40]:
        if p.style and p.style.name == "Title" and (p.text or "").strip():
            return norm_text((p.text or "").strip())
    return title


def _heading_level(p: Paragraph) -> int | None:
    if not p.style or not p.style.name.startswith("Heading"):
        return None
    try:
        return int(p.style.name.replace("Heading", "").strip())
    except ValueError:
        return None


def _md_heading_prefix(docx_heading_level: int) -> str:
    return "#" * (docx_heading_level + 1)


def table_to_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = " ".join(
                paragraph_to_md(pp) for pp in cell.paragraphs if paragraph_to_md(pp)
            ).strip()
            cells.append(text)
        rows.append(cells)
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    for r in rows:
        while len(r) < width:
            r.append("")
    headers = rows[0]
    body = rows[1:]
    lines = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * width) + "|",
    ]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def export_md(
    *,
    docx_path: Path,
    preserved_toc: str,
    abbrev_note: str,
    fallback_fig_paths: list[str],
) -> str:
    by_hash = _build_repo_hash_index()
    blip_paths = _resolve_blip_paths(docx_path, by_hash)
    blip_i = 0
    fig_fallback_i = 0

    def consume_figure_path() -> str | None:
        nonlocal blip_i, fig_fallback_i
        rel: str | None = None
        while blip_i < len(blip_paths) and rel is None:
            rel = blip_paths[blip_i]
            blip_i += 1
        if rel is None and fig_fallback_i < len(fallback_fig_paths):
            rel = fallback_fig_paths[fig_fallback_i]
            fig_fallback_i += 1
        return rel

    doc = Document(str(docx_path))
    title = _title_from_docx(doc)

    acknowledgements: list[str] = []
    abbrev_md: list[str] = []
    abstract: list[str] = []
    chapter_lines: list[str] = []

    phase = "find_ack"
    expect_fig_caption = False
    pending_figure_rel: str | None = None

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            p = block
            t_raw = (p.text or "").strip()
            t = norm_text(t_raw)
            st = p.style.name if p.style else ""
            hl = _heading_level(p)

            if phase == "find_ack":
                if hl == 1 and t == "Acknowledgements":
                    phase = "ack"
                continue

            if phase == "ack":
                if hl == 1 and t == "List of Abbreviations":
                    phase = "abbrev"
                elif t:
                    line = paragraph_to_md(p)
                    if line:
                        if st.startswith("List Bullet"):
                            acknowledgements.append(f"- {line}")
                        elif st.startswith("List Number"):
                            acknowledgements.append(line)
                        else:
                            acknowledgements.append(line)
                continue

            if phase == "abbrev":
                if hl == 1 and re.match(r"^1\.\s*Abstract", t):
                    phase = "abstract"
                continue

            if phase == "abstract":
                if hl == 1 and t == "Table of Contents":
                    phase = "skip_toc"
                elif hl is None and t:
                    abstract.append(paragraph_to_md(p))
                continue

            if phase == "skip_toc":
                if hl == 1 and t.startswith("Chapter "):
                    phase = "chapter"
                    chapter_lines.append(f"{_md_heading_prefix(1)} {norm_text(t)}")
                continue

            if phase == "chapter":
                if expect_fig_caption and st == "Figure Caption" and t:
                    cap_plain = norm_text(t)
                    cap_md = cap_plain if cap_plain.startswith("**") else f"**{cap_plain}**"
                    rel = pending_figure_rel
                    pending_figure_rel = None
                    if rel:
                        chapter_lines.append(f"![{cap_plain}]({rel})")
                    else:
                        chapter_lines.append("<!-- image: not mapped to repo asset -->")
                    chapter_lines.append(cap_md)
                    expect_fig_caption = False
                    continue
                expect_fig_caption = False

                if hl is not None and t:
                    chapter_lines.append(f"{_md_heading_prefix(hl)} {paragraph_to_md(p)}")
                elif st.startswith("List Bullet") and t:
                    chapter_lines.append(f"- {paragraph_to_md(p)}")
                elif st.startswith("List Number") and t:
                    chapter_lines.append(paragraph_to_md(p))
                elif t:
                    line = paragraph_to_md(p).strip()
                    if (
                        re.match(r"^\d+\.\d+\s+\S", line)
                        and not line.startswith("#")
                        and "http" not in line
                    ):
                        line = f"### {line}"
                    chapter_lines.append(line)
                continue

        elif isinstance(block, Table):
            if phase == "abbrev":
                abbrev_md.append(table_to_markdown(block))
                continue
            if phase == "chapter" and table_is_figure_frame(block):
                pending_figure_rel = consume_figure_path()
                expect_fig_caption = True
                continue
            if phase == "chapter":
                md_t = table_to_markdown(block)
                if md_t.strip():
                    chapter_lines.append(md_t)

    if expect_fig_caption and pending_figure_rel:
        chapter_lines.append(f"![]({pending_figure_rel})")
    elif expect_fig_caption:
        chapter_lines.append("<!-- image: caption missing in docx -->")

    meta_lines = [
        "**Arka Talukder | B01821011**  ",
        "**MSc Cyber Security (Full-time)**  ",
        "**University of the West of Scotland**  ",
        "**School of Computing, Engineering and Physical Sciences**  ",
        "**Supervisor: Dr. Raja Ujjan**",
    ]

    out: list[str] = [
        f"# {title}",
        "",
        *meta_lines,
        "",
        "---",
        "",
        "## 1. Abstract",
        "",
        *abstract,
        "",
        "---",
        "",
        "## Acknowledgements",
        "",
        *acknowledgements,
        "",
        "---",
        "",
        "## List of Abbreviations",
        "",
        *abbrev_md,
        "",
    ]
    if abbrev_note:
        out.append(abbrev_note)
        out.append("")
    out.extend(
        [
            "---",
            "",
        ]
    )
    if preserved_toc.strip():
        out.append(_strip_trailing_horizontal_rules(preserved_toc))
        out.append("")
    out.extend(
        [
            "---",
            "",
            *chapter_lines,
            "",
        ]
    )
    return "\n".join(out)


def main() -> None:
    global DOCX_PATH, MD_PATH
    p = argparse.ArgumentParser(description="Export Humanized (or DRAFT) .docx -> .md")
    p.add_argument("--dry-run", action="store_true", help="Print to stdout only")
    p.add_argument(
        "--docx",
        type=Path,
        help=f"Input Word file (default: {DEFAULT_DOCX.name})",
    )
    p.add_argument(
        "--out-md",
        type=Path,
        help=f"Output Markdown (default: {DEFAULT_MD.name})",
    )
    p.add_argument(
        "--toc-from",
        type=Path,
        help="Existing .md to preserve Table of Contents / LOF / LOT block from (default: out-md, else Humanized.md)",
    )
    args = p.parse_args()

    if args.docx is not None:
        docx_p = (args.docx if args.docx.is_absolute() else (ROOT / args.docx)).resolve()
    else:
        docx_p = DEFAULT_DOCX

    if args.out_md is not None:
        out_md = (args.out_md if args.out_md.is_absolute() else (ROOT / args.out_md)).resolve()
    else:
        out_md = (docx_p.with_suffix(".md") if args.docx is not None else DEFAULT_MD)

    if args.toc_from is not None:
        toc_path = (args.toc_from if args.toc_from.is_absolute() else (ROOT / args.toc_from)).resolve()
    elif out_md.is_file():
        toc_path = out_md
    else:
        toc_path = DEFAULT_MD

    DOCX_PATH = docx_p
    MD_PATH = out_md

    if not docx_p.is_file():
        print(f"ERROR: missing {docx_p}", file=sys.stderr)
        sys.exit(1)
    if not toc_path.is_file():
        print(
            f"ERROR: need a Markdown file for preserved TOC/LOF/LOT: {toc_path}",
            file=sys.stderr,
        )
        sys.exit(1)

    old = toc_path.read_text(encoding="utf-8")
    preserved = _extract_preserved_toc_block(old)
    note = _extract_abbrev_note(old)
    if not note:
        note = (
            "*Note: In the final Word file, abbreviations in the first column may be coloured "
            "(e.g. dark red) to match the programme’s preferred sample layout.*"
        )
    fallbacks = _figure_fallbacks_from_md(old)
    md = export_md(
        docx_path=DOCX_PATH,
        preserved_toc=preserved,
        abbrev_note=note,
        fallback_fig_paths=fallbacks,
    )

    if args.dry_run:
        print(md[:8000])
        print("\n... truncated ...\n")
        return
    MD_PATH.write_text(md, encoding="utf-8", newline="\n")
    print(f"Wrote {MD_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
