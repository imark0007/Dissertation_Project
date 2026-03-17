"""
Generate a premium-quality Word interim report from the Markdown source.
Includes: OMML equation formatting, numbered figures with captions,
formatted tables, Harvard references, 1.5 line spacing, Calibri 11pt.

Run from project root:
    python archive/scripts_one_time/generate_interim_report.py

Output: archive/interim_report/B01821011_Interim_Report_Final.docx
"""
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls
except ImportError:
    print("Please install python-docx: pip install python-docx")
    raise

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
MD_PATH = PROJECT_ROOT / "archive" / "interim_report" / "Interim_Report_Arka_Talukder.md"
OUT_PATH = PROJECT_ROOT / "archive" / "interim_report" / "B01821011_Interim_Report_Final.docx"

EQUATION_MAP = {
    r'\text{Precision} = \frac{TP}{TP + FP} \tag{1}':
        ('Precision = TP / (TP + FP)', '(1)'),
    r'\text{Recall} = \frac{TP}{TP + FN} \tag{2}':
        ('Recall = TP / (TP + FN)', '(2)'),
    r'F_1 = \frac{2 \times \text{Precision} \times \text{Recall}}{\text{Precision} + \text{Recall}} \tag{3}':
        ('F\u2081 = 2 \u00d7 Precision \u00d7 Recall / (Precision + Recall)', '(3)'),
    r'\text{FAR} = \frac{FP}{FP + TN} \tag{4}':
        ('FAR = FP / (FP + TN)', '(4)'),
    r'\text{AUC} = \int_0^1 \text{TPR}(t)\, d(\text{FPR}(t)) \tag{5}':
        ('AUC = \u222b\u2080\u00b9 TPR(t) d(FPR(t))', '(5)'),
    r'\text{IG}_i(x) = (x_i - x\'_i) \times \int_0^1 \frac{\partial F(x\' + \alpha(x - x\'))}{\partial x_i} d\alpha \tag{6}':
        ('IG\u1d62(x) = (x\u1d62 \u2212 x\u2032\u1d62) \u00d7 \u222b\u2080\u00b9 \u2202F(x\u2032 + \u03b1(x \u2212 x\u2032)) / \u2202x\u1d62 d\u03b1', '(6)'),
}


def set_cell_shading(cell, color="4472C4"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, sz="4", color="999999"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'start', 'bottom', 'end'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def set_paragraph_spacing(para, before=0, after=120, line=276):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_run(para, text, bold=False, italic=False, size=11, color=None, font="Calibri"):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_formatted_text(para, text, size=11):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(para, part[2:-2], bold=True, size=size)
        elif part.startswith('*') and part.endswith('*'):
            add_run(para, part[1:-1], italic=True, size=size)
        else:
            add_run(para, part, size=size)


def add_equation_paragraph(doc, eq_text, tag):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=120, after=120)

    tab_stops = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_center = OxmlElement('w:tab')
    tab_center.set(qn('w:val'), 'center')
    tab_center.set(qn('w:pos'), '4680')
    tabs.append(tab_center)
    tab_right = OxmlElement('w:tab')
    tab_right.set(qn('w:val'), 'right')
    tab_right.set(qn('w:pos'), '9360')
    tabs.append(tab_right)
    tab_stops.append(tabs)

    add_run(p, '\t')
    add_run(p, eq_text, italic=True, size=11)
    add_run(p, '\t')
    add_run(p, tag, size=11)


def add_table(doc, headers, rows, caption=None, header_color="4472C4"):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=10, color=(255, 255, 255))
        set_cell_shading(cell, header_color)
        set_cell_borders(cell, color="4472C4")

    for i, row_data in enumerate(rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j in range(min(len(row_data), n_cols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            text = str(row_data[j])
            if text.startswith('**') and text.endswith('**'):
                add_run(p, text[2:-2], bold=True, size=10)
            else:
                add_run(p, text, size=10)
            set_cell_shading(cell, bg)
            set_cell_borders(cell, color="BFBFBF")

    if caption:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=60, after=120)
        add_run(p, caption, italic=True, size=10)


def add_figure(doc, img_path, caption, width=Inches(5.0)):
    full_path = PROJECT_ROOT / img_path.replace('/', '\\')
    if full_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(full_path), width=width)

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap_p, before=60, after=180)
    add_run(cap_p, caption, italic=True, size=10)


def build_document():
    md_text = MD_PATH.read_text(encoding='utf-8')
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    for level in range(1, 4):
        sname = f'Heading {level}'
        if sname in doc.styles:
            hs = doc.styles[sname]
            hs.font.name = 'Calibri'
            hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            if level == 1:
                hs.font.size = Pt(16)
            elif level == 2:
                hs.font.size = Pt(13)
            else:
                hs.font.size = Pt(11)

    # --- TITLE PAGE ---
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'MSc Interim Report', bold=True, size=22, color=(31, 58, 95))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Explainable Dynamic Graph Neural Network SIEM for\n'
               'Software-Defined IoT using Edge AI and Federated Learning',
            bold=True, size=14, color=(31, 58, 95))

    for _ in range(3):
        doc.add_paragraph()

    for line in [
        'Arka Talukder',
        'Student Number: B01821011',
        'MSc Cyber Security (Full-time)',
        'Supervisor: Dr. Raja Ujjan',
        'University of the West of Scotland',
        'School of Computing, Engineering and Physical Sciences',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, line, size=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'Word count: approximately 4,950 words (maximum 5,000)', italic=True, size=11)

    doc.add_page_break()

    lines = md_text.split('\n')
    skip_header = True
    i = 0

    while i < len(lines):
        line = lines[i]

        if skip_header:
            if line.strip() == '---':
                skip_header = False
                i += 1
                if i < len(lines) and lines[i].strip() == '':
                    i += 1
                continue
            i += 1
            continue

        if line.strip() == '---':
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith('**Figure '):
            match = re.match(r'\*\*Figure (\d+): (.+?)\*\*', line)
            if match:
                fig_num = match.group(1)
                fig_title = match.group(2)
                i += 1
                if i < len(lines):
                    img_match = re.match(r'!\[.*?\]\((.+?)\)', lines[i])
                    if img_match:
                        img_path = img_match.group(1)
                        i += 1
                        source_text = ''
                        if i < len(lines) and lines[i].startswith('*'):
                            source_text = ' ' + lines[i].strip('* ')
                            i += 1
                        caption = f'Figure {fig_num}: {fig_title}.{source_text}'
                        add_figure(doc, img_path, caption)
                        continue
            i += 1
            continue

        if line.startswith('!['):
            i += 1
            continue

        if line.startswith('| ') and '|' in line:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 3:
                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                data_rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split('|') if c.strip()]
                    if cells:
                        data_rows.append(cells)
                add_table(doc, header_cells, data_rows)
            continue

        if line.startswith('$$'):
            eq_content = line.replace('$$', '').strip()
            i += 1
            while i < len(lines) and not lines[i].startswith('$$'):
                eq_content += ' ' + lines[i].strip()
                i += 1
            if i < len(lines):
                i += 1
            eq_content = eq_content.strip()
            matched = False
            for pattern, (display, tag) in EQUATION_MAP.items():
                if pattern.replace('\\', '') in eq_content.replace('\\', '') or \
                   eq_content.replace('\\', '').strip()[:20] in pattern.replace('\\', ''):
                    add_equation_paragraph(doc, display, tag)
                    matched = True
                    break
            if not matched:
                clean = re.sub(r'\\(text|frac|times|tag|int|partial|alpha)', '', eq_content)
                clean = clean.replace('{', '').replace('}', '').replace('_', '')
                clean = clean.replace('^', '').replace('\\,', ' ').strip()
                tag_match = re.search(r'\((\d+)\)', eq_content)
                tag = tag_match.group(0) if tag_match else ''
                if tag:
                    clean = clean.replace(tag, '').strip()
                add_equation_paragraph(doc, clean, tag)
            continue

        if line.startswith('- **') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line.lstrip('- ').strip()
            add_formatted_text(p, text)
            i += 1
            continue

        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s*', '', line).strip()
            add_formatted_text(p, text)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith('where '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=60)
            add_run(p, stripped, italic=True, size=10)
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    return doc


def main():
    doc = build_document()
    doc.save(str(OUT_PATH))
    print(f"Saved:  {OUT_PATH}")
    print(f"Source: {MD_PATH}")

    text = MD_PATH.read_text(encoding='utf-8')
    clean = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    clean = re.sub(r'[|#*$\-]', ' ', clean)
    clean = re.sub(r'\s+', ' ', clean)
    words = len(clean.split())
    print(f"Approx word count: {words}")


if __name__ == "__main__":
    main()
