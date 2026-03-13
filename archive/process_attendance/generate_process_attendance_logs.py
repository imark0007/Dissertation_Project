"""
Generate Project Process Documentation and Attendance Log for Arka Talukder.
Period: 6 February 2026 - 6 March 2026
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "project process and Attandance log")


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def create_process_documentation():
    doc = Document()
    
    # Header
    add_para(doc, "MSc PROJECT (COMP11024)", bold=True)
    add_para(doc, "PROJECT PROCESS DOCUMENTATION TEMPLATE", bold=True)
    doc.add_paragraph()
    
    add_para(doc, "Student: Arka Talukder (B01821011)              Supervisor: Dr. Raja Ujjan")
    add_para(doc, "Meeting Number: 1–6               Date/Time: 06/02/2026 – 06/03/2026")
    doc.add_paragraph()
    
    add_para(doc, "Agenda for meeting(s):", bold=True)
    doc.add_paragraph("• Project title selection and approval")
    doc.add_paragraph("• Project specification and scope discussion")
    doc.add_paragraph("• Interim report structure and guidelines")
    doc.add_paragraph("• Literature review and research gap alignment")
    doc.add_paragraph("• Methodology and experimental design review")
    doc.add_paragraph("• Progress on data pipeline and model implementation")
    doc.add_paragraph()
    
    add_para(doc, "Discussion of agenda items:", bold=True)
    discussion = """Across all meetings, the project moved from title selection to interim report completion. On 10 and 11 February 2026, the project title – Explainable Dynamic Graph Neural Network SIEM for Software-Defined IoT using Edge AI and Federated Learning – was discussed and approved by the supervisor. After approval, we clarified the scope. The project will use CICIoT2023 flow data, a dynamic GNN with GAT and GRU, federated learning with non-IID data, and SIEM alert generation.

The project specification was reviewed. The supervisor agreed with the four main parts: dynamic GNN, federated learning, explainability, and SIEM integration. We also discussed the interim report structure. It has three parts: literature review, methodology, and plan for completion.

I shared progress on the literature review. The supervisor gave feedback on the research gap table and how to compare my project with related work. We reviewed the methodology section. This included graph construction, model architectures (Random Forest, MLP, centralised GNN, federated GNN), and evaluation metrics.

I presented the implementation progress. This covered the data pipeline, graph builder, all four models, federated learning with Flower, explainability with Captum, and the FastAPI SIEM endpoint. The supervisor noted that the implementation matches the research objectives. The interim report draft was discussed. I will seek approval in the next meeting."""
    doc.add_paragraph(discussion)
    doc.add_paragraph()
    
    add_para(doc, "Summary of agreed action plan:", bold=True)
    doc.add_paragraph("• Obtain supervisor approval for the interim report in the next meeting")
    doc.add_paragraph("• Finalise any minor edits based on supervisor feedback")
    doc.add_paragraph("• Proceed with final experiments and results chapter for the dissertation")
    doc.add_paragraph("• Maintain documentation and reproducibility of all experiments")
    doc.add_paragraph()
    
    add_para(doc, "Notes:", bold=True)
    doc.add_paragraph("Supervisor emphasised clear alignment between research questions, methodology, and results. Careful documentation of experiments and fixed seeds for reproducibility was advised. The project is on track for timely completion.")
    
    out_path = os.path.join(OUTPUT_DIR, "Process_Documentation_B01821011_Arka_Talukder.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


def create_attendance_log():
    doc = Document()
    
    # Header - replicate sample structure
    add_para(doc, "MSc Dissertation Project Monthly Attendance and Progress Report", bold=True)
    doc.add_paragraph()
    doc.add_paragraph("TO BE COMPLETED BY THE STUDENT, WITH COMMENTS FROM THE SUPERVISOR.")
    doc.add_paragraph("Please complete all sections below")
    doc.add_paragraph()
    
    add_para(doc, "Name", bold=True)
    doc.add_paragraph("Arka Talukder")
    doc.add_paragraph()
    
    add_para(doc, "Student ID", bold=True)
    doc.add_paragraph("B01821011")
    doc.add_paragraph()
    
    add_para(doc, "Supervisor", bold=True)
    doc.add_paragraph("Dr. Raja Ujjan")
    doc.add_paragraph()
    
    add_para(doc, "Start Date", bold=True)
    doc.add_paragraph("06/02/2026")
    doc.add_paragraph()
    
    add_para(doc, "Mode of Study", bold=True)
    doc.add_paragraph("Online and Onsite")
    doc.add_paragraph()
    
    add_para(doc, "Dates of Meetings (include dates of all meetings)", bold=True)
    doc.add_paragraph("06/02/2026 (Online)")
    doc.add_paragraph("10/02/2026 (Onsite) – Title selection")
    doc.add_paragraph("11/02/2026 (Onsite) – Title approval")
    doc.add_paragraph("13/02/2026 (Online)")
    doc.add_paragraph("20/02/2026 (Online)")
    doc.add_paragraph("27/02/2026 (Online)")
    doc.add_paragraph("06/03/2026 (Online)")
    doc.add_paragraph()
    
    add_para(doc, "Description of Meetings", bold=True)
    desc = """During this reporting period, I held regular supervisory meetings with Dr. Raja Ujjan to discuss my dissertation project on explainable dynamic graph neural network SIEM for Software-Defined IoT.

In the first meetings (6 February), we discussed the project scope and what the interim report should include. On 10 and 11 February, the project title was selected and approved by the supervisor. After that, we focused on the project specification. This includes four main parts: dynamic GNN, federated learning, explainability, and SIEM integration.

Later meetings covered the literature review structure, research gap, and methodology design. The supervisor gave feedback on graph construction, model choices, and evaluation metrics. We also discussed how to structure the interim report and present the work clearly.

Recent meetings reviewed my implementation progress and the interim report draft. I presented the data pipeline, all four models (Random Forest, MLP, centralised GNN, federated GNN), explainability module, and SIEM API. The interim report is about 80% complete. I will seek final approval from the supervisor in the next week."""
    doc.add_paragraph(desc)
    doc.add_paragraph()
    
    add_para(doc, "Summary of work undertaken this month.", bold=True)
    work = """During this month, I completed most of the work needed for the interim report. Key achievements include:

• Project title approved (10–11 February 2026)
• Literature review written, covering IoT security, SDIoT, ML for intrusion detection, GNNs, federated learning, and explainable AI
• Research gap table completed, comparing my project with related work
• Methodology section written: dataset (CICIoT2023), preprocessing, graph construction, model architectures, evaluation metrics
• Data pipeline implemented: CICIoT2023 loaded, preprocessed, normalised, exported as Parquet
• Graph construction implemented: kNN builder, windowing, sequence dataset
• All four models implemented: Random Forest, MLP, centralised dynamic GNN (GAT+GRU), federated GNN
• Federated learning setup: Flower framework, non-IID data splitting, 10-round simulation
• Explainability: Integrated Gradients (Captum) and GAT attention weights
• SIEM integration: FastAPI endpoint producing ECS-formatted JSON alerts
• Plan for completion section written with timeline and contingency plans
• Interim report draft completed (approximately 80%); awaiting supervisor approval"""
    doc.add_paragraph(work)
    doc.add_paragraph()
    
    add_para(doc, "What work will you undertake next month?", bold=True)
    next_work = """In the upcoming period, I will focus on:

• Obtaining supervisor approval for the interim report
• Finalising any edits based on supervisor feedback
• Completing final experiments and generating all figures (confusion matrices, ROC curves, FL convergence)
• Writing the Results and Discussion chapters for the final dissertation
• Conducting sensitivity analysis if time permits
• Writing Conclusion and Critical Reflection
• Final proofreading and submission

I remain on schedule for timely completion."""
    doc.add_paragraph(next_work)
    doc.add_paragraph()
    
    add_para(doc, "Please detail reasons for any absence including the total number of days absent (annual leave, conference attendance, field research, ill health etc)?", bold=True)
    doc.add_paragraph("N/A")
    doc.add_paragraph()
    
    add_para(doc, "Statement from Supervisor (including any issues which should be brought to the attention of School, indication of satisfactory process thus far and whether or not attendance has been satisfactory)", bold=True)
    doc.add_paragraph("[To be completed by Dr. Raja Ujjan]")
    doc.add_paragraph()
    
    add_para(doc, "Signed (student)", bold=True)
    doc.add_paragraph("Date: 06/03/2026")
    doc.add_paragraph()
    
    add_para(doc, "Signed (MSc Dissertation Supervisor)", bold=True)
    doc.add_paragraph("Date:")
    
    out_path = os.path.join(OUTPUT_DIR, "Attendance_B01821011_Arka_Talukder_Feb-Mar.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    create_process_documentation()
    create_attendance_log()
    print("Done.")
