"""One-off generator: Mar–Apr 2026 attendance + process docs from school templates."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent

ATT_SRC = ROOT / "Arka Talukder_Attendance_Jan-Feb_B01821011.docx"
PROC_SRC = ROOT / "Arka_Talukder_Process_Documentation_B01821011.docx"

ATT_OUT = ROOT / "Arka_Talukder_Attendance_Mar-Apr_B01821011.docx"
PROC_OUT = ROOT / "Arka_Talukder_Process_Documentation_Mar-Apr_B01821011.docx"

MEETINGS_CELL = """07/03/2026 (Online) - future work: coding environment, results, references, implementation
14/03/2026 (Online) - same standing agenda (environment, results, references, implementation)
21/03/2026 (Online)
28/03/2026 (Online)
04/04/2026 (Online)
11/04/2026 (Online)
14/04/2026 (On campus) - full project review with supervisor
15/04/2026 - detailed written feedback and correction points received by email (following campus meeting); all items subsequently addressed"""

DESC_MEETINGS = """During this reporting period, I continued structured weekly supervision aligned to dissertation completion. On each Friday meeting (online), I reported progress on my coding environment, running and documenting experiments, compiling results, managing references, and implementation details for the Explainable dynamic GNN SIEM pipeline. On 14 April 2026 I met Dr. Raja Ujjan on campus for a full walkthrough of the project so that he could review the integrated system, dissertation structure, and evidence base in one session. On 15 April 2026 I received comprehensive feedback and a set of correction points by email. I have since worked through that list systematically and resolved the supervisor’s comments."""

SUMMARY_WORK = """During this month, I focused on turning the implementation into a submission-ready dissertation package. Key points include:
• Weekly Friday online meetings covering reproducible environment setup, experimental results, bibliography and citation hygiene, and implementation refinements
• On-campus supervision on 14 April 2026 with a complete project demonstration and discussion of the full dissertation draft
• Receipt of detailed email feedback on 15 April 2026; all requested corrections addressed (text, structure, figures/tables, and implementation notes as directed)
• Tightened alignment between research questions, methodology, results, and discussion in the final chapters"""

NEXT_MONTH = """In the upcoming period, I will focus on:
• Final internal consistency and proofreading pass across all chapters
• Final checks on figures, tables, and Harvard-style references
• Preparing final submission and supervisor handover materials
• Any short follow-up with the supervisor if a final sign-off meeting is needed
I remain on schedule for timely completion."""

ABSENCE = """No absence to report for scheduled supervision during this period."""

STUDENT_SIGN_DATE = "22/04/2026"


def fill_attendance() -> None:
    shutil.copy2(ATT_SRC, ATT_OUT)
    doc = Document(str(ATT_OUT))
    t0 = doc.tables[0]
    # Keep Name / ID / Supervisor / project Start Date; refresh meetings
    t0.rows[2].cells[3].text = MEETINGS_CELL
    if len(t0.rows) > 3:
        t0.rows[3].cells[3].text = ""

    doc.tables[1].rows[0].cells[1].text = DESC_MEETINGS
    doc.tables[2].rows[1].cells[0].text = SUMMARY_WORK
    doc.tables[3].rows[1].cells[0].text = NEXT_MONTH
    doc.tables[4].rows[1].cells[0].text = ABSENCE
    # Supervisor statement left blank
    doc.tables[6].rows[0].cells[3].text = STUDENT_SIGN_DATE

    doc.save(str(ATT_OUT))


def fill_process() -> None:
    shutil.copy2(PROC_SRC, PROC_OUT)
    doc = Document(str(PROC_OUT))

    doc.paragraphs[10].text = (
        "Meeting Number: 7-14\t\t\t\t\t\tDate/Time: 07/03/2026 - 22/04/2026"
    )

    bullets = [
        "Weekly Friday progress: coding environment, experimental results, references, and implementation",
        "Preparation for full on-campus project review with supervisor",
        "On-campus supervision: integrated demonstration and dissertation discussion (14 April 2026)",
        "Respond to written corrections received by email (15 April 2026)",
        "Final dissertation polish: structure, evidence, and reproducibility",
        "Cross-check chapters so objectives, methodology, results, and discussion stay aligned",
    ]
    for i, text in enumerate(bullets, start=15):
        doc.paragraphs[i].text = text

    doc.paragraphs[24].text = (
        "Across this block of meetings, supervision shifted from interim reporting to final "
        "dissertation delivery. Standing Friday sessions were used to keep the coding environment "
        "reproducible, consolidate experimental results, tighten references, and close implementation "
        "gaps in the dynamic GNN, federated learning, explainability, and SIEM integration tracks."
    )
    doc.paragraphs[26].text = (
        "The on-campus meeting on 14 April 2026 allowed the supervisor to review the whole project "
        "in context: end-to-end pipeline, model behaviour, evaluation outputs, and how these are "
        "presented in the dissertation. Discussion covered clarity of claims, traceability from "
        "objectives to evidence, and presentation of limitations."
    )
    doc.paragraphs[28].text = (
        "On 15 April 2026 I received detailed feedback and a structured list of corrections by "
        "email. This covered narrative edits, figure/table labelling, citation consistency, and "
        "targeted notes on methodology and results interpretation where alignment needed to be "
        "stronger."
    )
    doc.paragraphs[30].text = (
        "I have implemented all corrections requested. The dissertation draft now reflects the "
        "supervisor’s guidance from both the campus session and the follow-up email, with updated "
        "text, refreshed outputs where needed, and improved cross-referencing between implementation "
        "and discussion."
    )

    actions = [
        "Freeze main experiment configuration and document environment for reproducibility",
        "Retain evidence of applied corrections (tracked edits / version used for submission)",
        "Complete final proofread and reference audit before submission",
        "Confirm any last supervisor sign-off if required by School process",
    ]
    for i, text in enumerate(actions, start=38):
        doc.paragraphs[i].text = f"• {text}"

    doc.paragraphs[44].text = (
        "Supervisor feedback from the campus review plus the 15 April email provided a clear, "
        "actionable checklist. All points have been addressed. The project remains on track for "
        "final submission with documentation and scripts aligned to the reported results."
    )

    doc.save(str(PROC_OUT))


def main() -> None:
    fill_attendance()
    fill_process()
    print("Wrote:", ATT_OUT)
    print("Wrote:", PROC_OUT)


if __name__ == "__main__":
    main()
